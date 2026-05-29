from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"D:\DE_TAI")
OUTPUT = OUT / "PhanTich_ThietKe_DoAn_BanhKem_AI_Chinh_Sua_Theo_Mau.docx"

SCHOOL = "TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
TITLE = "XÂY DỰNG WEB BÁN BÁNH KEM TÍCH HỢP AI"
TOPIC = "Website bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"
CLASS_NAME = "22CT3"
COURSE = "DHCQ_K2022"
INTERNSHIP_UNIT = "Trung tâm Phát triển Phần mềm SDC – Đại học Đà Nẵng"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "EDEDED")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_para(doc, text="", bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(12)


def add_cover(doc):
    add_para(doc, SCHOOL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    p = add_para(doc, "ĐỀ CƯƠNG", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(20)
    p = add_para(doc, "THỰC TẬP TỐT NGHIỆP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(18)
    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_table(rows=7, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    rows = [
        ("Sinh viên thực hiện:", STUDENT),
        ("Mã sinh viên:", MSSV),
        ("Giáo viên hướng dẫn:", TEACHER),
        ("Lớp:", CLASS_NAME),
        ("Khóa:", COURSE),
        ("Tên đề tài:", TITLE),
        ("Đơn vị thực tập:", INTERNSHIP_UNIT),
    ]
    for row, (left, right) in zip(info.rows, rows):
        set_cell_text(row.cells[0], left, bold=True)
        set_cell_text(row.cells[1], right)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Đà Nẵng, năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "MỤC LỤC", level=1)
    entries = [
        "1. Tên đề tài",
        "2. Đơn vị thực tập",
        "3. Đặt vấn đề",
        "4. Mục tiêu đề tài",
        "5. Đối tượng sử dụng và phạm vi",
        "6. Chức năng hệ thống",
        "7. Công nghệ sử dụng và kiến trúc hệ thống",
        "8. Kế hoạch thực hiện",
    ]
    for entry in entries:
        add_para(doc, entry)
    doc.add_page_break()


def build_document():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. Tên đề tài", level=2)
    add_para(doc, TOPIC + ".")

    add_heading(doc, "2. Đơn vị thực tập", level=2)
    add_para(doc, f"- {INTERNSHIP_UNIT}")
    add_para(
        doc,
        "Đây là môi trường phù hợp để thực hiện đề tài phát triển ứng dụng web thương mại điện tử kết hợp trí tuệ nhân tạo, "
        "giúp sinh viên vận dụng kiến thức lập trình web, thiết kế cơ sở dữ liệu, xây dựng API và tích hợp dịch vụ AI vào một sản phẩm thực tế."
    )

    add_heading(doc, "3. Đặt vấn đề", level=2)
    add_para(
        doc,
        "Thị trường bánh kem thủ công tại Việt Nam ngày càng phát triển, đặc biệt với nhu cầu đặt bánh theo dịp sinh nhật, "
        "kỷ niệm, tiệc công ty và các sự kiện cá nhân. Tuy nhiên, nhiều tiệm bánh vẫn nhận đơn chủ yếu qua Facebook, Zalo hoặc gọi điện, "
        "khiến quá trình tư vấn, ghi chú yêu cầu, báo giá và theo dõi trạng thái đơn hàng còn phụ thuộc nhiều vào thao tác thủ công."
    )
    add_para(
        doc,
        "Khách hàng thường khó hình dung mẫu bánh trước khi đặt, dễ bỏ sót thông tin như kích thước, vị bánh, màu kem, topping, ngày nhận hoặc ghi chú đặc biệt. "
        "Về phía tiệm bánh, nhân viên phải đọc lại tin nhắn, tổng hợp yêu cầu và chuyển thông tin cho thợ làm bánh, từ đó có thể phát sinh sai sót trong quá trình sản xuất."
    )
    add_para(
        doc,
        "Vì vậy, đề tài xây dựng một website bán bánh kem tích hợp AI nhằm số hóa quy trình đặt bánh, hỗ trợ khách hàng tự thiết kế bánh, "
        "nhận tư vấn tự động, tạo phiếu đặt hàng rõ ràng cho thợ làm bánh và giúp chủ tiệm quản lý đơn hàng tập trung hơn."
    )

    add_heading(doc, "4. Mục tiêu đề tài", level=2)
    goals = [
        "Xây dựng website bán bánh kem có giao diện thân thiện, hỗ trợ khách hàng xem danh sách sản phẩm, tìm kiếm, lọc và xem chi tiết bánh.",
        "Phát triển chức năng thiết kế bánh Click-to-Customize cho phép khách hàng chọn kích thước, vị bánh, màu kem, topping và ghi chú yêu cầu.",
        "Tích hợp chatbot AI để tư vấn phong cách bánh theo dịp, kiểm tra thông tin còn thiếu và hỗ trợ khách hàng trước khi đặt.",
        "Tự động tính giá theo lựa chọn và tạo phiếu đặt hàng chi tiết gồm tóm tắt yêu cầu, ghi chú cho thợ và trạng thái đơn.",
        "Xây dựng khu vực quản trị cho Admin/Chủ tiệm để quản lý sản phẩm, đơn hàng, lịch sản xuất, thống kê doanh thu và phân quyền tài khoản.",
        "Hỗ trợ thợ làm bánh xem phiếu sản xuất, ghi chú và cập nhật trạng thái làm bánh.",
        "Tạo nền tảng có thể mở rộng thêm thanh toán trực tuyến, thông báo Zalo OA, quản lý tồn kho và ứng dụng mobile trong tương lai.",
    ]
    for goal in goals:
        add_bullet(doc, goal)

    add_heading(doc, "5. Đối tượng sử dụng và phạm vi", level=2)
    add_para(doc, "Hệ thống hướng đến hoạt động bán bánh kem trực tuyến cho một tiệm bánh hoặc cửa hàng nhỏ, tập trung vào quy trình khách hàng chọn bánh, đặt bánh, nhận tư vấn AI và theo dõi trạng thái đơn hàng.")
    add_para(doc, "Đối tượng sử dụng:")
    add_table(
        doc,
        ["Tác nhân", "Vai trò", "Chức năng chính"],
        [
            ("Khách vãng lai", "Người truy cập chưa đăng nhập", "Xem sản phẩm, tìm kiếm/lọc, xem chi tiết, đăng ký, đăng nhập, dùng chatbot tư vấn cơ bản."),
            ("Khách hàng", "Người dùng đã có tài khoản", "Thiết kế bánh, đặt bánh, chọn lịch nhận, thanh toán COD, theo dõi trạng thái đơn, đánh giá sản phẩm."),
            ("Admin / Chủ tiệm", "Người quản trị hệ thống", "Quản lý sản phẩm, đơn hàng, lịch sản xuất, doanh thu, tài khoản và phân quyền."),
            ("Thợ làm bánh", "Người tiếp nhận phiếu sản xuất", "Xem đơn cần làm, xem ghi chú sản xuất, cập nhật trạng thái đang làm/sẵn sàng."),
            ("Hệ thống AI", "Dịch vụ hỗ trợ tự động", "Tư vấn khách hàng, gợi ý sản phẩm, tạo tóm tắt đơn hàng, ghi chú cho thợ và lưu lịch sử trò chuyện."),
        ],
        widths=[1.45, 1.8, 3.9],
    )
    add_para(doc, "Phạm vi chức năng:")
    for item in [
        "Tập trung vào nền tảng Web App chạy trên desktop và thiết bị di động.",
        "Ưu tiên thanh toán khi nhận bánh (COD) để phù hợp phạm vi triển khai ban đầu.",
        "Quản lý vòng đời đơn hàng: pending, confirmed, in_production, ready, delivered.",
        "Chưa tập trung vào mobile app native, mô hình nhiều cửa hàng, quản lý kho nguyên liệu tự động hoặc huấn luyện mô hình AI riêng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Chức năng hệ thống", level=2)
    function_groups = [
        ("Quản lý tài khoản", "Đăng ký, đăng nhập, đăng xuất, xác thực người dùng, quản lý thông tin tài khoản và phân quyền theo vai trò."),
        ("Xem và tìm kiếm sản phẩm", "Xem danh sách bánh, xem chi tiết sản phẩm, tìm kiếm theo tên, lọc theo loại bánh, kích thước, mức giá hoặc dịp sử dụng."),
        ("Thiết kế bánh Click-to-Customize", "Tùy chỉnh kích thước, vị bánh, màu kem, topping, ghi chú yêu cầu và xem giá tự động theo lựa chọn."),
        ("Chatbot AI tư vấn", "Tư vấn phong cách bánh theo dịp, gợi ý mẫu phù hợp, hỏi lại thông tin còn thiếu và lưu lịch sử trò chuyện."),
        ("Đặt bánh và thanh toán", "Tạo đơn hàng, chọn ngày nhận, nhập thông tin nhận bánh, xác nhận đơn và chọn thanh toán khi nhận bánh (COD)."),
        ("Tạo phiếu đặt hàng bằng AI", "AI tổng hợp yêu cầu thành phiếu rõ ràng gồm tóm tắt đơn, ghi chú cho thợ, giá dự kiến và hạn hoàn thành."),
        ("Quản lý đơn hàng", "Admin xem danh sách đơn, lọc/tìm đơn, xem chi tiết, xác nhận đơn, cập nhật đã giao và xử lý yêu cầu bổ sung/hủy đơn."),
        ("Quản lý sản xuất", "Thợ làm bánh xem đơn cần làm, xem phiếu sản xuất, ghi chú và cập nhật trạng thái in_production hoặc ready."),
        ("Thống kê và báo cáo", "Admin xem thống kê doanh thu, số lượng đơn, sản phẩm bán chạy và xuất báo cáo phục vụ quản lý."),
    ]
    for name, detail in function_groups:
        add_para(doc, f"{name}: {detail}")

    add_heading(doc, "7. Công nghệ sử dụng và kiến trúc hệ thống", level=2)
    add_para(
        doc,
        "Hệ thống được xây dựng theo mô hình Client–Server kết hợp kiến trúc 3 tầng và một lớp dịch vụ AI. "
        "Cách tổ chức này giúp giao diện, xử lý nghiệp vụ, dữ liệu và AI được tách riêng, dễ bảo trì và dễ mở rộng."
    )
    add_table(
        doc,
        ["Thành phần", "Công nghệ", "Chức năng"],
        [
            ("Frontend", "Next.js 14, React, TypeScript, Tailwind CSS", "Hiển thị giao diện người dùng, Cake Builder, trang sản phẩm, giỏ/đơn hàng và khu vực quản trị."),
            ("Backend", "FastAPI, Python", "Xử lý nghiệp vụ, xác thực, quản lý sản phẩm, đơn hàng, sản xuất và cung cấp RESTful API."),
            ("Database", "Supabase PostgreSQL", "Lưu trữ người dùng, sản phẩm, đơn hàng, tùy chỉnh bánh, lịch sử chat và trạng thái sản xuất."),
            ("Storage/Auth", "Supabase Auth, Supabase Storage", "Quản lý tài khoản, phân quyền, lưu ảnh sản phẩm và tài nguyên hệ thống."),
            ("AI Services", "Claude API / Anthropic, Prompt Engineering, RAG", "Tư vấn khách hàng, sinh phiếu đặt hàng, kiểm tra thông tin còn thiếu và tạo ghi chú cho thợ."),
            ("Triển khai", "Vercel, Docker", "Triển khai frontend, backend và chuẩn bị môi trường demo/staging."),
            ("Công cụ hỗ trợ", "GitHub, VS Code, Postman, Figma", "Quản lý mã nguồn, kiểm thử API, thiết kế giao diện và theo dõi tiến độ."),
        ],
        widths=[1.45, 2.25, 3.45],
    )
    add_para(doc, "Kiến trúc tổng quát:")
    for item in [
        "Tầng giao diện: Next.js hiển thị website bán bánh, trang thiết kế bánh và dashboard quản trị.",
        "Tầng xử lý nghiệp vụ: FastAPI cung cấp API cho sản phẩm, đơn hàng, tài khoản, sản xuất và AI.",
        "Tầng dữ liệu: Supabase PostgreSQL lưu trữ dữ liệu có cấu trúc, Supabase Storage lưu ảnh sản phẩm.",
        "Lớp AI: AIService gọi Claude API, xử lý prompt và trả về kết quả tư vấn hoặc phiếu đặt hàng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Kế hoạch thực hiện", level=2)
    add_table(
        doc,
        ["Thời gian", "Nội dung công việc", "Sản phẩm đầu ra"],
        [
            (
                "Tuần 1: 18/05 - 24/05",
                "Khảo sát nhu cầu đặt bánh online; phân tích quy trình bán qua mạng xã hội; xác định actor, use case và phạm vi hệ thống.",
                "Tài liệu khảo sát; danh sách chức năng; Use Case tổng quan.",
            ),
            (
                "Tuần 2: 25/05 - 31/05",
                "Thiết kế kiến trúc hệ thống; phân tích cơ sở dữ liệu; thiết kế quy trình đặt bánh và activity diagram; xây dựng prototype giao diện.",
                "Sơ đồ kiến trúc; thiết kế database; Activity Diagram; prototype UI.",
            ),
            (
                "Tuần 3: 01/06 - 07/06",
                "Xây dựng frontend với Next.js; phát triển trang sản phẩm, tìm kiếm/lọc, chi tiết sản phẩm và giao diện Cake Builder.",
                "Giao diện website; chức năng xem/tìm sản phẩm; bản đầu của Cake Builder.",
            ),
            (
                "Tuần 4: 08/06 - 14/06",
                "Xây dựng backend FastAPI; phát triển API tài khoản, sản phẩm, đơn hàng, trạng thái đơn và khu vực quản trị.",
                "API backend; chức năng đặt hàng; dashboard quản trị cơ bản.",
            ),
            (
                "Tuần 5: 15/06 - 21/06",
                "Tích hợp AI chatbot; xây dựng module tạo phiếu đặt hàng; kiểm thử quy trình đặt bánh, xác nhận đơn và cập nhật sản xuất.",
                "Chatbot AI; phiếu đặt hàng tự động; quy trình end-to-end hoạt động.",
            ),
            (
                "Tuần 6: 22/06 - 28/06",
                "Hoàn thiện giao diện; tối ưu hiệu năng; kiểm thử chức năng; chuẩn bị báo cáo, slide thuyết trình và demo hệ thống.",
                "Bản demo hoàn chỉnh; báo cáo tổng kết; slide thuyết trình.",
            ),
        ],
        widths=[1.35, 3.6, 2.2],
    )

    try:
        doc.save(OUTPUT)
        return OUTPUT
    except PermissionError:
        fallback = OUTPUT.with_name(OUTPUT.stem + "_Can_Le.docx")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build_document())
