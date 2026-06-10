import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_env_out = os.environ.get("DE_TAI_OUTPUT_DIR")
if _env_out:
    OUT = Path(_env_out)
else:
    # Fall back to project root (parent of tools/) or home dir
    _tools_dir = Path(__file__).resolve().parent
    _project_root = _tools_dir.parent
    OUT = _project_root if (_project_root / "backend").exists() else Path.home()

DOCX_PATH = OUT / "Bao_Cao_Tien_Do_Thuc_Tap_Tuan_2_Ly_Thanh_Long.docx"

SCHOOL = "TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
TITLE = "BÁO CÁO TIẾN ĐỘ THỰC TẬP TUẦN 2"
TOPIC = "Xây dựng Web bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"
WEEK = "Tuần 2: 25/05/2026 - 31/05/2026"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, color="000000")
    return paragraph


def add_para(doc, text="", bold=False, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 2"].font.name = "Times New Roman"


def add_info_table(doc):
    rows = [
        ("Sinh viên thực hiện", STUDENT),
        ("Mã sinh viên", MSSV),
        ("Giáo viên hướng dẫn", TEACHER),
        ("Đề tài", TOPIC),
        ("Tuần báo cáo", WEEK),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        set_cell_shading(table.rows[i].cells[0], "EDEDED")
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_doc(doc)

    add_para(doc, SCHOOL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    title = add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title.runs[0], size=18, bold=True)
    doc.add_paragraph()
    add_info_table(doc)

    add_heading(doc, "1. Công việc đã thực hiện trong tuần", level=1)
    for item in [
        "Hoàn thiện lại tài liệu phân tích thiết kế theo bố cục mẫu, gồm tên đề tài, đặt vấn đề, mục tiêu, phạm vi, chức năng hệ thống, công nghệ và kế hoạch thực hiện.",
        "Rà soát và chỉnh sửa phần căn lề, định dạng trang Word theo mẫu báo cáo đã tham khảo.",
        "Thiết kế kiến trúc tổng thể của hệ thống theo mô hình Client-Server kết hợp 3 tầng: giao diện, xử lý nghiệp vụ, dữ liệu và lớp dịch vụ AI.",
        "Xác định các bảng dữ liệu chính dự kiến sử dụng: users, products, orders, cake_customizations, chat_history và production_notes.",
        "Hoàn thiện lại Use Case Diagram theo hướng tách từng tác nhân riêng để dễ đọc và dễ báo cáo.",
        "Chỉnh sửa Activity Diagram cho quy trình khách hàng đặt bánh và thanh toán COD, đảm bảo có đúng một điểm kết thúc.",
        "Xác định các màn hình giao diện chính cần thiết cho prototype UI: trang chủ, danh sách sản phẩm, chi tiết bánh, thiết kế bánh, chatbot AI, đặt hàng, theo dõi đơn, admin và thợ làm bánh.",
        "Tìm hiểu hướng dùng công cụ AI như Stitch/Figma Make/Antigravity để hỗ trợ tạo prototype UI ban đầu.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Kết quả đạt được", level=1)
    for item in [
        "Có bản tài liệu phân tích thiết kế được trình bày lại gọn gàng, theo cấu trúc tương tự mẫu tham khảo.",
        "Xác định rõ kiến trúc hệ thống và công nghệ dự kiến: Next.js, Tailwind CSS, FastAPI, Supabase/PostgreSQL và Claude API.",
        "Có danh sách bảng dữ liệu chính phục vụ bước thiết kế cơ sở dữ liệu chi tiết.",
        "Có bộ sơ đồ Use Case tách theo từng tác nhân: Khách vãng lai, Khách hàng, Admin/Chủ tiệm, Thợ làm bánh và Hệ thống AI.",
        "Có Activity Diagram hoàn chỉnh cho quy trình đặt bánh, thể hiện luồng xử lý, điều kiện rẽ nhánh, fork/join và điểm kết thúc.",
        "Có danh sách màn hình cần thiết để triển khai prototype UI trong tuần tiếp theo.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Khó khăn/vướng mắc", level=1)
    for item in [
        "Cần cân đối phạm vi đề tài vì hệ thống có nhiều nhóm chức năng: bán hàng, quản trị, sản xuất và AI.",
        "Việc trình bày Use Case Diagram ban đầu còn nhiều đường nối, cần tách theo tác nhân để tránh rối.",
        "Activity Diagram phải chỉnh nhiều lần để đúng ký hiệu UML và chỉ giữ một điểm kết thúc.",
        "Cần lựa chọn cách làm prototype UI phù hợp, vừa nhanh để báo cáo tiến độ vừa có thể chỉnh sửa tiếp trên Figma.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Kế hoạch thực hiện tuần tiếp theo", level=1)
    for item in [
        "Thiết kế chi tiết cơ sở dữ liệu, xác định trường dữ liệu, khóa chính, khóa ngoại và quan hệ giữa các bảng.",
        "Xây dựng prototype UI cho các màn hình chính bằng Stitch/Figma Make hoặc Figma.",
        "Bắt đầu dựng frontend với Next.js và Tailwind CSS cho các màn hình: trang chủ, danh sách sản phẩm, chi tiết sản phẩm và thiết kế bánh.",
        "Chuẩn bị cấu trúc backend FastAPI cho các module sản phẩm, tài khoản và đơn hàng.",
        "Xây dựng dữ liệu mẫu cho sản phẩm bánh, lựa chọn tùy chỉnh và trạng thái đơn hàng.",
        "Tiếp tục cập nhật tài liệu nếu giảng viên góp ý thêm trong buổi báo cáo.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign.add_run("Sinh viên thực hiện\n\n\nLý Thành Long")
    set_run_font(run, bold=True)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except (PermissionError, FileNotFoundError, OSError):
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        try:
            doc.save(fallback)
            return fallback
        except (PermissionError, FileNotFoundError, OSError):
            last_resort = Path.cwd() / DOCX_PATH.name
            doc.save(last_resort)
            return last_resort


if __name__ == "__main__":
    print(build_docx())
