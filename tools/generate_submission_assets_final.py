# -*- coding: utf-8 -*-
import os
import sys
from pathlib import Path
from textwrap import wrap

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import Circle, Ellipse, FancyArrowPatch, FancyBboxPatch, Polygon, Rectangle

# Configurable output directory
_out_env = os.environ.get("ASSETS_OUT_DIR")
if not _out_env:
    for idx, arg in enumerate(sys.argv):
        if arg == "--out-dir" and idx + 1 < len(sys.argv):
            _out_env = sys.argv[idx + 1]
            break

if _out_env:
    OUT = Path(_out_env)
else:
    OUT = Path(__file__).resolve().parent.parent.parent / "tai_lieu_nop"

USECASE_PATH = OUT / "so_do_use_case_banh_kem_ai.png"
ACTIVITY_PATH = OUT / "so_do_flowchart_dat_banh_ai.png"
DOCX_PATH = OUT / "Mo_ta_chuc_nang_Web_Ban_Banh_Kem_Tich_Hop_AI.docx"

TITLE = "XÂY DỰNG WEB BÁN BÁNH KEM TÍCH HỢP AI"
TOPIC = "Website bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"

PINK = "#E8837A"
CREAM = "#FDF6EE"
MOCHA = "#5C3D2E"
BLUE = "#B9D7F0"
YELLOW = "#FFE7A8"
MINT = "#BEE7D2"
LAVENDER = "#D9CCFF"
GRAY = "#665D58"
BG = "#FFFDF9"

plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False


def wrapped_text(text: str, width: int = 20) -> str:
    return "\n".join(wrap(text, width=width, break_long_words=False))


def add_actor(ax, x, y, name):
    ax.add_patch(Circle((x, y + 0.36), 0.12, facecolor="white", edgecolor=MOCHA, linewidth=1.4))
    ax.plot([x, x], [y + 0.24, y - 0.22], color=MOCHA, linewidth=1.4)
    ax.plot([x - 0.24, x + 0.24], [y + 0.04, y + 0.04], color=MOCHA, linewidth=1.4)
    ax.plot([x, x - 0.22], [y - 0.22, y - 0.52], color=MOCHA, linewidth=1.4)
    ax.plot([x, x + 0.22], [y - 0.22, y - 0.52], color=MOCHA, linewidth=1.4)
    ax.text(x, y - 0.78, wrapped_text(name, 13), ha="center", va="top", fontsize=8.6, color=MOCHA, fontweight="bold")


def add_ellipse(ax, x, y, w, h, text, fc="white", fontsize=7.2, width=16):
    ellipse = Ellipse((x, y), w, h, facecolor=fc, edgecolor=PINK, linewidth=1.35)
    ax.add_patch(ellipse)
    ax.text(x, y, wrapped_text(text, width), ha="center", va="center", fontsize=fontsize, color=MOCHA)
    return ellipse


def connect(ax, start, end, dashed=False, lw=0.8, alpha=0.82):
    ax.plot(
        [start[0], end[0]],
        [start[1], end[1]],
        color=GRAY,
        linewidth=lw,
        linestyle="--" if dashed else "-",
        alpha=alpha,
    )


def create_usecase_diagram():
    fig, ax = plt.subplots(figsize=(17, 10.8), dpi=200)
    ax.set_xlim(0, 17)
    ax.set_ylim(0, 10.8)
    ax.axis("off")
    ax.set_facecolor(BG)

    ax.add_patch(Rectangle((2.2, 0.85), 12.6, 8.85, facecolor=CREAM, edgecolor=MOCHA, linewidth=1.8))
    ax.text(8.5, 10.18, "Use Case Diagram - Hệ thống Web bán bánh kem tích hợp AI", ha="center", va="center", fontsize=16, color=MOCHA, fontweight="bold")

    add_actor(ax, 0.85, 8.35, "Khách vãng lai")
    add_actor(ax, 0.85, 5.15, "Khách hàng")
    add_actor(ax, 16.1, 7.55, "Admin / Chủ tiệm")
    add_actor(ax, 16.1, 3.45, "Thợ làm bánh")
    add_actor(ax, 16.1, 1.35, "Hệ thống AI")

    ax.text(5.6, 9.42, "Khách / mua hàng", ha="center", va="center", fontsize=8.5, color=MOCHA, fontweight="bold")
    ax.text(11.95, 9.42, "Quản trị & sản xuất", ha="center", va="center", fontsize=8.5, color=MOCHA, fontweight="bold")
    ax.text(6.25, 2.35, "Dịch vụ AI", ha="center", va="center", fontsize=8.5, color=MOCHA, fontweight="bold")

    use_cases = {
        # Customer and visitor
        "register": (3.55, 8.82, "Đăng ký tài khoản", "white"),
        "login": (5.55, 8.82, "Đăng nhập", "white"),
        "list": (3.55, 8.02, "Xem danh sách sản phẩm", "white"),
        "search": (5.55, 8.02, "Tìm kiếm / lọc sản phẩm", "white"),
        "detail": (7.55, 8.02, "Xem chi tiết sản phẩm", "white"),
        "logout": (3.55, 6.78, "Đăng xuất", "white"),
        "design": (5.55, 6.78, "Thiết kế bánh Click-to-Customize", "white"),
        "choose": (7.55, 6.78, "Chọn kích thước, vị, kem, màu, topping", "white"),
        "price": (9.55, 6.78, "Xem giá tự động", BLUE),
        "chat": (3.55, 5.78, "Chatbot AI tư vấn", BLUE),
        "order": (5.55, 5.78, "Đặt bánh", "white"),
        "pickup": (7.55, 5.78, "Chọn lịch nhận bánh", "white"),
        "cod": (9.55, 5.78, "Chọn thanh toán khi nhận bánh (COD)", "white"),
        "track": (4.55, 4.78, "Theo dõi trạng thái đơn", "white"),
        "review": (6.95, 4.78, "Đánh giá sản phẩm", "white"),
        # Admin
        "admin_list_product": (10.85, 8.72, "Xem danh sách sản phẩm quản trị", YELLOW),
        "admin_add_product": (12.85, 8.72, "Thêm sản phẩm", YELLOW),
        "admin_edit_product": (10.85, 7.92, "Sửa sản phẩm", YELLOW),
        "admin_upload_image": (12.85, 7.92, "Upload ảnh sản phẩm", YELLOW),
        "admin_toggle_product": (11.85, 7.12, "Ẩn / hiện sản phẩm", YELLOW),
        "admin_list_order": (10.85, 5.95, "Xem danh sách đơn hàng", YELLOW),
        "admin_filter_order": (12.85, 5.95, "Lọc / tìm đơn hàng", YELLOW),
        "admin_order_detail": (10.85, 5.15, "Xem chi tiết đơn hàng", YELLOW),
        "admin_confirm": (12.85, 5.15, "Xác nhận đơn hàng", YELLOW),
        "admin_delivered": (11.85, 4.35, "Đánh dấu đã giao", YELLOW),
        # Baker
        "baker_orders": (10.85, 3.45, "Xem đơn cần làm", MINT),
        "baker_ticket": (12.85, 3.45, "Xem phiếu làm bánh", MINT),
        "baker_notes": (10.85, 2.65, "Ghi chú sản xuất", MINT),
        "baker_status": (12.85, 2.65, "Cập nhật trạng thái đang làm / sẵn sàng", MINT),
        # AI
        "ai_question": (4.25, 1.55, "Xử lý câu hỏi tư vấn", LAVENDER),
        "ai_recommend": (6.25, 1.55, "Gợi ý sản phẩm phù hợp", LAVENDER),
        "ai_summary": (8.25, 1.55, "Tạo tóm tắt / phiếu đặt hàng", LAVENDER),
        "ai_history": (6.25, 0.9, "Lưu lịch sử trò chuyện", LAVENDER),
    }

    for key, (x, y, label, fc) in use_cases.items():
        add_ellipse(ax, x, y, 1.72, 0.5, label, fc=fc, fontsize=6.25, width=15)

    # Actor associations. Lines are intentionally light and mostly short/fanned by area.
    visitor_targets = ["register", "login", "list", "search", "detail"]
    customer_targets = ["logout", "design", "choose", "price", "chat", "order", "pickup", "cod", "track", "review"]
    admin_targets = [
        "admin_list_product",
        "admin_add_product",
        "admin_edit_product",
        "admin_upload_image",
        "admin_toggle_product",
        "admin_list_order",
        "admin_filter_order",
        "admin_order_detail",
        "admin_confirm",
        "admin_delivered",
    ]
    baker_targets = ["baker_orders", "baker_ticket", "baker_notes", "baker_status"]
    ai_targets = ["ai_question", "ai_recommend", "ai_summary", "ai_history"]

    for t in visitor_targets:
        x, y, *_ = use_cases[t]
        connect(ax, (1.16, 8.35), (x - 0.86, y), lw=0.52, alpha=0.55)
    for t in customer_targets:
        x, y, *_ = use_cases[t]
        connect(ax, (1.16, 5.15), (x - 0.86, y), lw=0.52, alpha=0.55)
    for t in admin_targets:
        x, y, *_ = use_cases[t]
        connect(ax, (15.83, 7.55), (x + 0.86, y), lw=0.52, alpha=0.55)
    for t in baker_targets:
        x, y, *_ = use_cases[t]
        connect(ax, (15.83, 3.45), (x + 0.86, y), lw=0.52, alpha=0.55)
    for t in ai_targets:
        x, y, *_ = use_cases[t]
        connect(ax, (15.83, 1.35), (x + 0.86, y), dashed=True, lw=0.52, alpha=0.55)

    # Essential include relations only.
    include_pairs = [
        ("design", "choose"),
        ("design", "price"),
        ("order", "pickup"),
        ("order", "cod"),
        ("chat", "ai_question"),
        ("ai_question", "ai_recommend"),
        ("order", "ai_summary"),
    ]
    for a, b in include_pairs:
        x1, y1, *_ = use_cases[a]
        x2, y2, *_ = use_cases[b]
        connect(ax, (x1, y1 - 0.26), (x2, y2 + 0.26), dashed=True, lw=0.62, alpha=0.65)
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.08, "<<include>>", fontsize=5.6, color=GRAY, ha="center")

    ax.text(8.5, 0.35, f"Đề tài: {TITLE} | Thành viên: {STUDENT} ({MSSV})", ha="center", va="center", fontsize=8.5, color=MOCHA)
    fig.tight_layout(pad=0.35)
    fig.savefig(USECASE_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def add_activity(ax, x, y, w, h, text, fc="white", fontsize=7.6, width=23):
    patch = FancyBboxPatch(
        (x - w / 2, y - h / 2),
        w,
        h,
        boxstyle="round,pad=0.06,rounding_size=0.12",
        facecolor=fc,
        edgecolor=PINK,
        linewidth=1.35,
    )
    ax.add_patch(patch)
    ax.text(x, y, wrapped_text(text, width), ha="center", va="center", fontsize=fontsize, color=MOCHA)
    return patch


def add_decision(ax, x, y, w, h, text):
    points = [(x, y + h / 2), (x + w / 2, y), (x, y - h / 2), (x - w / 2, y)]
    patch = Polygon(points, closed=True, facecolor="#FFF3D5", edgecolor=YELLOW, linewidth=1.45)
    ax.add_patch(patch)
    ax.text(x, y, wrapped_text(text, 16), ha="center", va="center", fontsize=7.0, color=MOCHA)
    return patch


def add_start_end(ax, x, y, text, fc=MOCHA):
    ax.add_patch(Ellipse((x, y), 2.4, 0.42, facecolor=fc, edgecolor=fc, linewidth=1.2))
    ax.text(x, y, text, ha="center", va="center", fontsize=7.5, color="white", fontweight="bold")


def add_initial_node(ax, x, y):
    ax.add_patch(Circle((x, y), 0.13, facecolor="black", edgecolor="black", linewidth=1.0, zorder=5))


def add_final_node(ax, x, y):
    ax.add_patch(Circle((x, y), 0.18, facecolor="white", edgecolor="black", linewidth=1.2, zorder=5))
    ax.add_patch(Circle((x, y), 0.11, facecolor="black", edgecolor="black", linewidth=1.0, zorder=6))


def add_bar(ax, x, y, w=1.0):
    ax.add_patch(Rectangle((x - w / 2, y - 0.035), w, 0.07, facecolor="black", edgecolor="black", linewidth=1.0, zorder=4))


def arrow(ax, start, end, label=None, curve=0.0):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.25,
            color=GRAY,
            connectionstyle=f"arc3,rad={curve}",
        )
    )
    if label:
        ax.text((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 + 0.12, label, fontsize=7, color=GRAY, ha="center")


def _create_activity_diagram_v1():
    fig, ax = plt.subplots(figsize=(14, 18.5), dpi=200)
    ax.set_xlim(0, 15.8)
    ax.set_ylim(0, 18.5)
    ax.axis("off")
    ax.set_facecolor(BG)
    ax.text(7.9, 18.05, "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD", ha="center", va="center", fontsize=16, color=MOCHA, fontweight="bold")
    ax.text(7.9, 17.65, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=9, color=GRAY)

    lane_left = 0.45
    lane_bottom = 0.25
    lane_top = 17.25
    lane_h = lane_top - lane_bottom
    lane_w = [3.0, 3.0, 2.65, 2.65, 3.25]
    lane_titles = ["Khách hàng", "Hệ thống Web", "Hệ thống AI", "Admin / Chủ tiệm", "Thợ làm bánh"]
    lane_colors = ["#FFF7EF", "#FFFFFF", "#F6F0FF", "#FFF4CC", "#ECFFF5"]
    xs = [lane_left]
    for w in lane_w:
        xs.append(xs[-1] + w)
    centers = [(xs[i] + xs[i + 1]) / 2 for i in range(len(lane_w))]

    for i, title in enumerate(lane_titles):
        ax.add_patch(Rectangle((xs[i], lane_bottom), lane_w[i], lane_h, facecolor=lane_colors[i], edgecolor=GRAY, linewidth=0.9, linestyle="--"))
        ax.add_patch(Rectangle((xs[i], lane_top - 0.55), lane_w[i], 0.55, facecolor="#FFE7A8", edgecolor=GRAY, linewidth=0.9, linestyle="--"))
        ax.text(centers[i], lane_top - 0.28, title, ha="center", va="center", fontsize=8.2, color=MOCHA, fontweight="bold")

    C, W, AI, AD, B = centers
    node_w = 2.25
    node_h = 0.46

    add_initial_node(ax, C, 16.45)
    add_activity(ax, C, 15.85, node_w, node_h, "Truy cập website", "white", fontsize=6.6, width=17)
    add_activity(ax, W, 15.15, node_w, node_h, "Hiển thị menu và Cake Builder", "white", fontsize=6.6, width=17)
    add_activity(ax, C, 14.45, node_w, node_h, "Xem/lọc sản phẩm hoặc tùy chỉnh bánh", "white", fontsize=6.5, width=16)
    add_decision(ax, C, 13.65, 1.52, 0.72, "Cần AI tư vấn?")
    add_activity(ax, AI, 13.65, node_w, node_h, "Xử lý câu hỏi tư vấn", BLUE, fontsize=6.5, width=16)
    add_activity(ax, AI, 12.95, node_w, node_h, "Gợi ý bánh và lưu lịch sử", BLUE, fontsize=6.5, width=16)
    add_decision(ax, C, 12.95, 0.72, 0.45, "")
    add_activity(ax, W, 12.25, node_w, node_h, "Tính giá tự động và hiển thị tóm tắt", BLUE, fontsize=6.4, width=17)
    add_decision(ax, C, 11.5, 1.55, 0.72, "Xác nhận đặt bánh?")
    add_decision(ax, W, 10.75, 1.45, 0.70, "Đã đăng nhập?")
    add_activity(ax, C, 10.75, node_w, node_h, "Đăng ký / đăng nhập", "white", fontsize=6.5, width=16)
    add_decision(ax, W, 10.05, 0.72, 0.45, "")
    add_activity(ax, C, 9.35, node_w, node_h, "Nhập thông tin nhận bánh và ngày nhận", "white", fontsize=6.4, width=16)
    add_decision(ax, W, 8.62, 1.58, 0.72, "Ngày nhận hợp lệ?")
    add_activity(ax, C, 8.62, node_w, node_h, "Nhập lại ngày nhận", "#FFEDE9", fontsize=6.5, width=16)
    add_activity(ax, C, 7.85, node_w, node_h, "Chọn thanh toán COD", "white", fontsize=6.5, width=16)
    add_activity(ax, W, 7.18, node_w, node_h, "Tạo đơn pending và lưu tùy chỉnh", BLUE, fontsize=6.4, width=16)
    add_bar(ax, W, 6.62, w=1.1)
    add_activity(ax, W, 6.02, node_w, node_h, "Lưu lịch sử trạng thái đơn", BLUE, fontsize=6.4, width=16)
    add_activity(ax, AI, 6.02, node_w, node_h, "Tạo phiếu đặt hàng và ghi chú cho thợ", BLUE, fontsize=6.2, width=16)
    add_bar(ax, (W + AI) / 2, 5.42, w=3.45)
    add_activity(ax, AD, 4.88, node_w, node_h, "Xem chi tiết đơn hàng", YELLOW, fontsize=6.5, width=16)
    add_decision(ax, AD, 4.18, 1.52, 0.72, "Xác nhận đơn?")
    add_activity(ax, W, 4.18, node_w, node_h, "Thông báo bổ sung hoặc hủy đơn", "#FFEDE9", fontsize=6.3, width=16)
    add_activity(ax, AD, 3.48, node_w, node_h, "Cập nhật pending → confirmed", YELLOW, fontsize=6.4, width=16)
    add_activity(ax, B, 2.78, node_w, node_h, "Xem phiếu làm bánh", MINT, fontsize=6.5, width=16)
    add_activity(ax, B, 2.12, node_w, node_h, "Cập nhật in_production", MINT, fontsize=6.5, width=16)
    add_activity(ax, B, 1.46, node_w, node_h, "Cập nhật ready", MINT, fontsize=6.5, width=16)
    add_activity(ax, C, 1.46, node_w, node_h, "Nhận bánh và thanh toán COD", "white", fontsize=6.4, width=16)
    add_activity(ax, AD, 0.9, node_w, node_h, "Cập nhật delivered", YELLOW, fontsize=6.5, width=16)
    add_activity(ax, C, 0.9, node_w, node_h, "Đánh giá sản phẩm", "white", fontsize=6.5, width=16)
    add_final_node(ax, C, 0.43)

    arrow(ax, (C, 16.3), (C, 16.08))
    arrow(ax, (C, 15.62), (W, 15.38))
    arrow(ax, (W, 14.92), (C, 14.68))
    arrow(ax, (C, 14.22), (C, 14.02))
    arrow(ax, (C + 0.76, 13.65), (AI - 1.13, 13.65), "[Có]")
    arrow(ax, (AI, 13.42), (AI, 13.18))
    arrow(ax, (AI - 1.13, 12.95), (C + 0.36, 12.95))
    arrow(ax, (C, 13.29), (C, 13.18), "[Không]")
    arrow(ax, (C, 12.72), (W, 12.47))
    arrow(ax, (W, 12.02), (C, 11.86))
    arrow(ax, (C + 0.78, 11.5), (C + 1.15, 14.22), "[Không]", curve=0.26)
    arrow(ax, (C, 11.14), (W, 11.08), "[Có]")
    arrow(ax, (W - 0.72, 10.75), (C + 1.13, 10.75), "[Chưa]")
    arrow(ax, (C + 1.13, 10.55), (W - 0.36, 10.05))
    arrow(ax, (W, 10.4), (W, 10.28), "[Rồi]")
    arrow(ax, (W - 0.36, 10.05), (C, 9.58))
    arrow(ax, (C + 1.13, 9.35), (W - 0.78, 8.62))
    arrow(ax, (W - 0.78, 8.62), (C + 1.13, 8.62), "[Không]")
    arrow(ax, (C, 8.39), (C, 9.12), curve=-0.2)
    arrow(ax, (W, 8.26), (C, 8.08), "[Có]")
    arrow(ax, (C + 1.13, 7.85), (W - 1.13, 7.18))
    arrow(ax, (W, 6.95), (W, 6.66))
    arrow(ax, (W, 6.58), (W, 6.25))
    arrow(ax, (W + 0.32, 6.58), (AI - 0.45, 6.25))
    arrow(ax, (W, 5.79), ((W + AI) / 2 - 0.6, 5.46))
    arrow(ax, (AI, 5.79), ((W + AI) / 2 + 0.6, 5.46))
    arrow(ax, ((W + AI) / 2, 5.38), (AD, 5.11))
    arrow(ax, (AD, 4.65), (AD, 4.54))
    arrow(ax, (AD - 0.76, 4.18), (W + 1.13, 4.18), "[Không]")
    arrow(ax, (W, 3.95), (C, 9.12), curve=0.32)
    arrow(ax, (AD, 3.82), (AD, 3.71), "[Có]")
    arrow(ax, (AD, 3.25), (B, 3.01))
    arrow(ax, (B, 2.55), (B, 2.35))
    arrow(ax, (B, 1.89), (B, 1.69))
    arrow(ax, (B - 1.13, 1.46), (C + 1.13, 1.46))
    arrow(ax, (C + 1.13, 1.24), (AD - 1.13, 0.9))
    arrow(ax, (AD - 1.13, 0.9), (C + 1.13, 0.9))
    arrow(ax, (C, 0.67), (C, 0.57))

    ax.text(13.45, 0.62, "Trạng thái đơn:\npending → confirmed →\nin_production → ready → delivered", ha="center", va="center", fontsize=7.0, color=MOCHA, bbox=dict(boxstyle="round,pad=0.25", fc="#FFF3D5", ec=YELLOW, lw=1.0))

    fig.tight_layout(pad=0.35)
    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def _create_activity_diagram_v2():
    fig, ax = plt.subplots(figsize=(14, 18), dpi=200)
    ax.set_xlim(0, 15.8)
    ax.set_ylim(0, 18)
    ax.axis("off")
    ax.set_facecolor(BG)
    ax.text(7.9, 17.55, "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD", ha="center", va="center", fontsize=16, color=MOCHA, fontweight="bold")
    ax.text(7.9, 17.17, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=9, color=GRAY)

    lane_left = 0.45
    lane_bottom = 0.35
    lane_top = 16.75
    lane_w = [3.05, 3.05, 2.7, 2.75, 3.25]
    lane_titles = ["Khách hàng", "Hệ thống Web", "Hệ thống AI", "Admin / Chủ tiệm", "Thợ làm bánh"]
    lane_colors = ["#FFF7EF", "#FFFFFF", "#F6F0FF", "#FFF4CC", "#ECFFF5"]
    xs = [lane_left]
    for w in lane_w:
        xs.append(xs[-1] + w)
    centers = [(xs[i] + xs[i + 1]) / 2 for i in range(len(lane_w))]

    for i, title in enumerate(lane_titles):
        ax.add_patch(Rectangle((xs[i], lane_bottom), lane_w[i], lane_top - lane_bottom, facecolor=lane_colors[i], edgecolor=GRAY, linewidth=0.9, linestyle="--"))
        ax.add_patch(Rectangle((xs[i], lane_top - 0.55), lane_w[i], 0.55, facecolor="#FFE7A8", edgecolor=GRAY, linewidth=0.9, linestyle="--"))
        ax.text(centers[i], lane_top - 0.28, title, ha="center", va="center", fontsize=8.3, color=MOCHA, fontweight="bold")

    C, W, AI, AD, B = centers
    node_w = 2.25
    node_h = 0.46

    # Nodes
    add_initial_node(ax, C, 15.95)
    add_activity(ax, C, 15.35, node_w, node_h, "Truy cập website", "white", fontsize=6.6, width=17)
    add_activity(ax, W, 14.7, node_w, node_h, "Hiển thị menu và Cake Builder", "white", fontsize=6.6, width=17)
    add_activity(ax, C, 14.05, node_w, node_h, "Xem/lọc sản phẩm hoặc tùy chỉnh bánh", "white", fontsize=6.5, width=16)
    add_decision(ax, C, 13.25, 1.5, 0.72, "Cần AI tư vấn?")
    add_activity(ax, AI, 13.25, node_w, node_h, "Tư vấn và gợi ý bánh phù hợp", BLUE, fontsize=6.4, width=16)
    add_activity(ax, W, 12.45, node_w, node_h, "Tính giá tự động và hiển thị tóm tắt", BLUE, fontsize=6.4, width=16)
    add_decision(ax, C, 11.65, 1.55, 0.72, "Xác nhận đặt bánh?")
    add_decision(ax, W, 10.85, 1.45, 0.70, "Đã đăng nhập?")
    add_activity(ax, C, 10.85, node_w, node_h, "Đăng ký / đăng nhập", "white", fontsize=6.5, width=16)
    add_activity(ax, C, 9.95, node_w, node_h, "Nhập thông tin nhận bánh và ngày nhận", "white", fontsize=6.4, width=16)
    add_decision(ax, W, 9.15, 1.58, 0.72, "Ngày nhận hợp lệ?")
    add_activity(ax, C, 9.15, node_w, node_h, "Nhập lại ngày nhận", "#FFEDE9", fontsize=6.5, width=16)
    add_activity(ax, C, 8.35, node_w, node_h, "Chọn thanh toán COD", "white", fontsize=6.5, width=16)
    add_activity(ax, W, 7.65, node_w, node_h, "Tạo đơn pending và lưu tùy chỉnh", BLUE, fontsize=6.4, width=16)
    add_activity(ax, AI, 6.95, node_w, node_h, "Tạo phiếu đặt hàng và ghi chú cho thợ", BLUE, fontsize=6.25, width=16)
    add_activity(ax, AD, 6.25, node_w, node_h, "Xem chi tiết đơn hàng", YELLOW, fontsize=6.5, width=16)
    add_decision(ax, AD, 5.45, 1.52, 0.72, "Đơn hợp lệ?")
    add_activity(ax, W, 5.45, node_w, node_h, "Thông báo bổ sung hoặc hủy đơn", "#FFEDE9", fontsize=6.3, width=16)
    add_activity(ax, AD, 4.65, node_w, node_h, "Xác nhận đơn: pending → confirmed", YELLOW, fontsize=6.25, width=16)
    add_activity(ax, B, 3.95, node_w, node_h, "Xem phiếu làm bánh", MINT, fontsize=6.5, width=16)
    add_activity(ax, B, 3.25, node_w, node_h, "Cập nhật in_production", MINT, fontsize=6.5, width=16)
    add_activity(ax, B, 2.55, node_w, node_h, "Cập nhật ready", MINT, fontsize=6.5, width=16)
    add_activity(ax, C, 1.85, node_w, node_h, "Nhận bánh và thanh toán COD", "white", fontsize=6.4, width=16)
    add_activity(ax, AD, 1.25, node_w, node_h, "Cập nhật delivered", YELLOW, fontsize=6.5, width=16)
    add_activity(ax, C, 0.8, node_w, node_h, "Đánh giá sản phẩm", "white", fontsize=6.5, width=16)
    add_final_node(ax, C, 0.42)

    # Flow
    arrow(ax, (C, 15.8), (C, 15.58))
    arrow(ax, (C, 15.12), (W, 14.93))
    arrow(ax, (W, 14.47), (C, 14.28))
    arrow(ax, (C, 13.82), (C, 13.61))
    arrow(ax, (C + 0.75, 13.25), (AI - 1.13, 13.25), "[Có]")
    arrow(ax, (AI, 13.02), (W + 1.13, 12.45))
    arrow(ax, (C, 12.89), (W - 1.13, 12.45), "[Không]")
    arrow(ax, (W, 12.22), (C, 12.01))
    arrow(ax, (C + 0.78, 11.65), (C + 1.05, 13.9), "[Không]", curve=0.22)
    arrow(ax, (C, 11.29), (W, 11.2), "[Có]")
    arrow(ax, (W - 0.72, 10.85), (C + 1.13, 10.85), "[Chưa]")
    arrow(ax, (C, 10.62), (C, 10.18))
    arrow(ax, (W, 10.5), (C, 10.18), "[Rồi]")
    arrow(ax, (C + 1.13, 9.95), (W - 0.79, 9.15))
    arrow(ax, (W - 0.79, 9.15), (C + 1.13, 9.15), "[Không]")
    arrow(ax, (C, 8.92), (C, 9.72), curve=-0.2)
    arrow(ax, (W, 8.79), (C, 8.58), "[Có]")
    arrow(ax, (C + 1.13, 8.35), (W - 1.13, 7.65))
    arrow(ax, (W + 1.13, 7.65), (AI - 1.13, 6.95))
    arrow(ax, (AI + 1.13, 6.95), (AD - 1.13, 6.25))
    arrow(ax, (AD, 6.02), (AD, 5.81))
    arrow(ax, (AD - 0.76, 5.45), (W + 1.13, 5.45), "[Không]")
    arrow(ax, (W, 5.22), (C, 9.72), curve=0.25)
    arrow(ax, (AD, 5.09), (AD, 4.88), "[Có]")
    arrow(ax, (AD + 1.13, 4.65), (B - 1.13, 3.95))
    arrow(ax, (B, 3.72), (B, 3.48))
    arrow(ax, (B, 3.02), (B, 2.78))
    arrow(ax, (B - 1.13, 2.55), (C + 1.13, 1.85))
    arrow(ax, (C + 1.13, 1.85), (AD - 1.13, 1.25))
    arrow(ax, (AD - 1.13, 1.25), (C + 1.13, 0.8))
    arrow(ax, (C, 0.57), (C, 0.54))

    ax.text(13.5, 0.7, "Trạng thái đơn:\npending → confirmed →\nin_production → ready → delivered", ha="center", va="center", fontsize=7.0, color=MOCHA, bbox=dict(boxstyle="round,pad=0.25", fc="#FFF3D5", ec=YELLOW, lw=1.0))

    fig.tight_layout(pad=0.35)
    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def _create_activity_diagram_v3():
    fig, ax = plt.subplots(figsize=(11, 16.5), dpi=200)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 16.5)
    ax.axis("off")
    ax.set_facecolor(BG)

    ax.text(
        5.5,
        16.12,
        "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD",
        ha="center",
        va="center",
        fontsize=14,
        color=MOCHA,
        fontweight="bold",
    )
    ax.text(5.5, 15.78, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=8, color=GRAY)

    def role_box(x, y, label, color):
        ax.add_patch(
            FancyBboxPatch(
                (x - 0.72, y - 0.14),
                1.44,
                0.28,
                boxstyle="round,pad=0.03,rounding_size=0.06",
                facecolor=color,
                edgecolor=GRAY,
                linewidth=0.7,
            )
        )
        ax.text(x, y, label, ha="center", va="center", fontsize=6.3, color=MOCHA)

    role_box(2.0, 15.35, "Khách hàng", "#FFF7EF")
    role_box(3.75, 15.35, "Hệ thống Web", "#EAF4FE")
    role_box(5.5, 15.35, "Hệ thống AI", "#EDE7FF")
    role_box(7.25, 15.35, "Admin", "#FFF3D5")
    role_box(9.0, 15.35, "Thợ làm bánh", "#E7F7EF")

    add_initial_node(ax, 5.5, 14.88)
    add_activity(ax, 5.5, 14.28, 5.6, 0.46, "Khách hàng: Truy cập website", "#FFF7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 13.58, 5.6, 0.46, "Hệ thống Web: Hiển thị menu và Cake Builder", "#EAF4FE", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 12.88, 5.6, 0.46, "Khách hàng: Xem/lọc sản phẩm hoặc tùy chỉnh bánh", "#FFF7EF", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 12.05, 2.05, 0.75, "Cần AI tư vấn?")
    add_activity(ax, 8.25, 12.05, 3.6, 0.46, "Hệ thống AI: Tư vấn và gợi ý bánh", "#EDE7FF", fontsize=6.8, width=24)
    add_decision(ax, 5.5, 11.25, 0.82, 0.48, "")
    add_activity(ax, 5.5, 10.55, 5.6, 0.46, "Hệ thống Web: Tính giá tự động và hiển thị tóm tắt", "#EAF4FE", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 9.75, 2.15, 0.75, "Khách xác nhận đặt bánh?")
    add_activity(ax, 2.05, 9.75, 2.9, 0.46, "Quay lại chỉnh sửa lựa chọn", "#FFF7EF", fontsize=6.8, width=20)
    add_decision(ax, 5.5, 8.9, 1.85, 0.72, "Đã đăng nhập?")
    add_activity(ax, 2.05, 8.9, 2.9, 0.46, "Đăng ký / đăng nhập", "#FFF7EF", fontsize=6.8, width=20)
    add_decision(ax, 5.5, 8.15, 0.82, 0.48, "")
    add_activity(ax, 5.5, 7.5, 5.6, 0.46, "Khách hàng: Nhập thông tin nhận bánh và ngày nhận", "#FFF7EF", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 6.72, 2.15, 0.75, "Ngày nhận hợp lệ?")
    add_activity(ax, 2.05, 6.72, 2.9, 0.46, "Nhập lại ngày nhận", "#FFEDE9", fontsize=6.8, width=20)
    add_activity(ax, 5.5, 5.92, 5.6, 0.46, "Khách hàng: Chọn thanh toán khi nhận bánh (COD)", "#FFF7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 5.22, 5.6, 0.46, "Hệ thống Web: Tạo đơn pending và lưu tùy chỉnh", "#EAF4FE", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 4.52, 5.6, 0.46, "Hệ thống AI: Tạo phiếu đặt hàng và ghi chú cho thợ", "#EDE7FF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 3.82, 5.6, 0.46, "Admin: Xem chi tiết đơn hàng", "#FFF3D5", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 3.05, 1.95, 0.72, "Đơn hợp lệ?")
    add_activity(ax, 2.05, 3.05, 2.9, 0.46, "Thông báo bổ sung hoặc hủy đơn", "#FFEDE9", fontsize=6.6, width=20)
    add_activity(ax, 5.5, 2.25, 5.6, 0.46, "Admin: Xác nhận đơn pending → confirmed", "#FFF3D5", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 1.58, 5.6, 0.46, "Thợ làm bánh: Xem phiếu, cập nhật in_production → ready", "#E7F7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 0.92, 5.6, 0.46, "Khách nhận bánh, thanh toán COD; Admin cập nhật delivered", "#FFF7EF", fontsize=6.9, width=35)
    add_activity(ax, 5.5, 0.38, 5.6, 0.38, "Khách hàng: Đánh giá sản phẩm", "#FFF7EF", fontsize=6.8, width=35)
    add_final_node(ax, 5.5, -0.02)

    arrow(ax, (5.5, 14.73), (5.5, 14.51))
    arrow(ax, (5.5, 14.05), (5.5, 13.81))
    arrow(ax, (5.5, 13.35), (5.5, 13.11))
    arrow(ax, (5.5, 12.65), (5.5, 12.43))
    arrow(ax, (6.52, 12.05), (6.45, 12.05), "[Có]")
    arrow(ax, (6.45, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    arrow(ax, (6.53, 12.05), (6.45, 12.05))
    # Actual branch arrows kept after repeated no-op anchors above to keep Mermaid-style layout separate from image layout.
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))

    # Redraw clean branch on top.
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    ax.add_patch(FancyArrowPatch((6.52, 12.05), (6.45, 12.05), arrowstyle="-|>", mutation_scale=0, linewidth=0, color=GRAY))
    arrow(ax, (6.52, 12.05), (6.48, 12.05))
    arrow(ax, (6.48, 12.05), (6.45, 12.05))
    # Direct custom horizontal arrow to the AI action.
    ax.add_patch(FancyArrowPatch((6.52, 12.05), (6.45, 12.05), arrowstyle="-", mutation_scale=0, linewidth=0, color=GRAY))
    ax.add_patch(FancyArrowPatch((6.52, 12.05), (6.45, 12.05), arrowstyle="-", mutation_scale=0, linewidth=0, color=GRAY))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))

    # Clearer branch and merge lines.
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    arrow(ax, (6.52, 12.05), (6.45, 12.05))
    # Matplotlib label positions for decision branches.
    ax.text(6.7, 12.24, "[Có]", fontsize=6.4, color=GRAY)
    ax.add_patch(FancyArrowPatch((6.52, 12.05), (6.45, 12.05), arrowstyle="-", mutation_scale=0, linewidth=0, color=GRAY))
    ax.add_patch(FancyArrowPatch((6.52, 12.05), (7.0, 12.05), arrowstyle="-|>", mutation_scale=10, linewidth=1.05, color=GRAY))
    arrow(ax, (8.25, 11.82), (5.88, 11.25))
    arrow(ax, (5.5, 11.69), (5.5, 11.49), "[Không]")
    arrow(ax, (5.5, 11.0), (5.5, 10.78))
    arrow(ax, (5.5, 10.32), (5.5, 10.11))
    arrow(ax, (4.43, 9.75), (3.5, 9.75), "[Không]")
    arrow(ax, (2.05, 9.52), (4.45, 12.65), curve=-0.24)
    arrow(ax, (5.5, 9.39), (5.5, 9.25), "[Có]")
    arrow(ax, (4.58, 8.9), (3.5, 8.9), "[Chưa]")
    arrow(ax, (2.05, 8.67), (5.08, 8.15))
    arrow(ax, (5.5, 8.54), (5.5, 8.39), "[Rồi]")
    arrow(ax, (5.5, 7.91), (5.5, 7.73))
    arrow(ax, (5.5, 7.27), (5.5, 7.1))
    arrow(ax, (4.43, 6.72), (3.5, 6.72), "[Không]")
    arrow(ax, (2.05, 6.49), (4.32, 7.5), curve=-0.18)
    arrow(ax, (5.5, 6.36), (5.5, 6.15), "[Có]")
    arrow(ax, (5.5, 5.69), (5.5, 5.45))
    arrow(ax, (5.5, 4.99), (5.5, 4.75))
    arrow(ax, (5.5, 4.29), (5.5, 4.05))
    arrow(ax, (5.5, 3.59), (5.5, 3.41))
    arrow(ax, (4.53, 3.05), (3.5, 3.05), "[Không]")
    arrow(ax, (2.05, 2.82), (4.35, 7.5), curve=-0.24)
    arrow(ax, (5.5, 2.69), (5.5, 2.48), "[Có]")
    arrow(ax, (5.5, 2.02), (5.5, 1.81))
    arrow(ax, (5.5, 1.35), (5.5, 1.15))
    arrow(ax, (5.5, 0.69), (5.5, 0.57))
    arrow(ax, (5.5, 0.19), (5.5, 0.12))

    ax.text(
        8.8,
        0.65,
        "Trạng thái đơn:\npending → confirmed →\nin_production → ready → delivered",
        ha="center",
        va="center",
        fontsize=6.6,
        color=MOCHA,
        bbox=dict(boxstyle="round,pad=0.25", fc="#FFF3D5", ec=YELLOW, lw=1.0),
    )

    fig.tight_layout(pad=0.4)
    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def _create_activity_diagram_v4():
    fig, ax = plt.subplots(figsize=(11, 16.5), dpi=200)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 16.5)
    ax.axis("off")
    ax.set_facecolor(BG)

    ax.text(
        5.5,
        16.12,
        "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD",
        ha="center",
        va="center",
        fontsize=14,
        color=MOCHA,
        fontweight="bold",
    )
    ax.text(5.5, 15.78, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=8, color=GRAY)

    legend = [
        (1.8, "Khách hàng", "#FFF7EF"),
        (3.55, "Hệ thống Web", "#EAF4FE"),
        (5.45, "Hệ thống AI", "#EDE7FF"),
        (7.25, "Admin", "#FFF3D5"),
        (9.0, "Thợ làm bánh", "#E7F7EF"),
    ]
    for x, label, color in legend:
        ax.add_patch(
            FancyBboxPatch(
                (x - 0.72, 15.20),
                1.44,
                0.3,
                boxstyle="round,pad=0.03,rounding_size=0.06",
                facecolor=color,
                edgecolor=GRAY,
                linewidth=0.7,
            )
        )
        ax.text(x, 15.35, label, ha="center", va="center", fontsize=6.3, color=MOCHA)

    add_initial_node(ax, 5.5, 14.85)
    add_activity(ax, 5.5, 14.25, 5.6, 0.46, "Khách hàng: Truy cập website", "#FFF7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 13.55, 5.6, 0.46, "Hệ thống Web: Hiển thị menu và Cake Builder", "#EAF4FE", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 12.85, 5.6, 0.46, "Khách hàng: Xem/lọc sản phẩm hoặc tùy chỉnh bánh", "#FFF7EF", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 12.05, 2.05, 0.75, "Cần AI tư vấn?")
    add_activity(ax, 8.4, 12.05, 3.2, 0.46, "Hệ thống AI: Tư vấn và gợi ý bánh", "#EDE7FF", fontsize=6.8, width=22)
    add_decision(ax, 5.5, 11.25, 0.82, 0.48, "")
    add_activity(ax, 5.5, 10.55, 5.6, 0.46, "Hệ thống Web: Tính giá tự động và hiển thị tóm tắt", "#EAF4FE", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 9.75, 2.15, 0.75, "Khách xác nhận đặt bánh?")
    add_activity(ax, 2.05, 9.75, 2.9, 0.46, "Quay lại chỉnh sửa lựa chọn", "#FFF7EF", fontsize=6.8, width=20)
    add_decision(ax, 5.5, 8.9, 1.85, 0.72, "Đã đăng nhập?")
    add_activity(ax, 2.05, 8.9, 2.9, 0.46, "Đăng ký / đăng nhập", "#FFF7EF", fontsize=6.8, width=20)
    add_decision(ax, 5.5, 8.15, 0.82, 0.48, "")
    add_activity(ax, 5.5, 7.5, 5.6, 0.46, "Khách hàng: Nhập thông tin nhận bánh và ngày nhận", "#FFF7EF", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 6.72, 2.15, 0.75, "Ngày nhận hợp lệ?")
    add_activity(ax, 2.05, 6.72, 2.9, 0.46, "Nhập lại ngày nhận", "#FFEDE9", fontsize=6.8, width=20)
    add_activity(ax, 5.5, 5.92, 5.6, 0.46, "Khách hàng: Chọn thanh toán khi nhận bánh (COD)", "#FFF7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 5.22, 5.6, 0.46, "Hệ thống Web: Tạo đơn pending và lưu tùy chỉnh", "#EAF4FE", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 4.52, 5.6, 0.46, "Hệ thống AI: Tạo phiếu đặt hàng và ghi chú cho thợ", "#EDE7FF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 3.82, 5.6, 0.46, "Admin: Xem chi tiết đơn hàng", "#FFF3D5", fontsize=7.0, width=35)
    add_decision(ax, 5.5, 3.05, 1.95, 0.72, "Đơn hợp lệ?")
    add_activity(ax, 2.05, 3.05, 2.9, 0.46, "Thông báo bổ sung hoặc hủy đơn", "#FFEDE9", fontsize=6.6, width=20)
    add_activity(ax, 5.5, 2.25, 5.6, 0.46, "Admin: Xác nhận đơn pending → confirmed", "#FFF3D5", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 1.55, 5.6, 0.46, "Thợ làm bánh: Xem phiếu, cập nhật in_production → ready", "#E7F7EF", fontsize=7.0, width=35)
    add_activity(ax, 5.5, 0.88, 5.6, 0.46, "Khách nhận bánh, thanh toán COD; Admin cập nhật delivered", "#FFF7EF", fontsize=6.9, width=35)
    add_activity(ax, 5.5, 0.38, 5.6, 0.36, "Khách hàng: Đánh giá sản phẩm", "#FFF7EF", fontsize=6.8, width=35)
    add_final_node(ax, 5.5, 0.05)

    arrow(ax, (5.5, 14.70), (5.5, 14.48))
    arrow(ax, (5.5, 14.02), (5.5, 13.78))
    arrow(ax, (5.5, 13.32), (5.5, 13.08))
    arrow(ax, (5.5, 12.62), (5.5, 12.43))
    arrow(ax, (6.53, 12.05), (6.8, 12.05), "[Có]")
    arrow(ax, (8.4, 11.82), (5.9, 11.25))
    arrow(ax, (5.5, 11.68), (5.5, 11.49), "[Không]")
    arrow(ax, (5.5, 11.01), (5.5, 10.78))
    arrow(ax, (5.5, 10.32), (5.5, 10.11))
    arrow(ax, (4.42, 9.75), (3.5, 9.75), "[Không]")
    arrow(ax, (2.05, 9.52), (4.45, 12.62), curve=-0.24)
    arrow(ax, (5.5, 9.38), (5.5, 9.25), "[Có]")
    arrow(ax, (4.57, 8.9), (3.5, 8.9), "[Chưa]")
    arrow(ax, (2.05, 8.67), (5.08, 8.15))
    arrow(ax, (5.5, 8.54), (5.5, 8.39), "[Rồi]")
    arrow(ax, (5.5, 7.91), (5.5, 7.73))
    arrow(ax, (5.5, 7.27), (5.5, 7.10))
    arrow(ax, (4.43, 6.72), (3.5, 6.72), "[Không]")
    arrow(ax, (2.05, 6.49), (4.32, 7.5), curve=-0.18)
    arrow(ax, (5.5, 6.35), (5.5, 6.15), "[Có]")
    arrow(ax, (5.5, 5.69), (5.5, 5.45))
    arrow(ax, (5.5, 4.99), (5.5, 4.75))
    arrow(ax, (5.5, 4.29), (5.5, 4.05))
    arrow(ax, (5.5, 3.59), (5.5, 3.41))
    arrow(ax, (4.53, 3.05), (3.5, 3.05), "[Không]")
    arrow(ax, (2.05, 2.82), (4.35, 7.5), curve=-0.24)
    arrow(ax, (5.5, 2.69), (5.5, 2.48), "[Có]")
    arrow(ax, (5.5, 2.02), (5.5, 1.78))
    arrow(ax, (5.5, 1.32), (5.5, 1.11))
    arrow(ax, (5.5, 0.65), (5.5, 0.56))
    arrow(ax, (5.5, 0.20), (5.5, 0.17))

    ax.text(
        8.8,
        0.65,
        "Trạng thái đơn:\npending → confirmed →\nin_production → ready → delivered",
        ha="center",
        va="center",
        fontsize=6.6,
        color=MOCHA,
        bbox=dict(boxstyle="round,pad=0.25", fc="#FFF3D5", ec=YELLOW, lw=1.0),
    )

    fig.tight_layout(pad=0.4)
    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def _create_activity_diagram_v5():
    fig, ax = plt.subplots(figsize=(11, 17.2), dpi=200)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 17.2)
    ax.axis("off")
    ax.set_facecolor(BG)

    ax.text(
        5.5,
        16.75,
        "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD",
        ha="center",
        va="center",
        fontsize=14,
        color=MOCHA,
        fontweight="bold",
    )
    ax.text(5.5, 16.38, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=8, color=GRAY)

    legend = [
        (1.55, "Khách hàng", "#FFF7EF"),
        (3.35, "Hệ thống Web", "#EAF4FE"),
        (5.45, "Hệ thống AI", "#EDE7FF"),
        (7.35, "Admin", "#FFF3D5"),
        (9.1, "Thợ làm bánh", "#E7F7EF"),
    ]
    for x, label, color in legend:
        ax.add_patch(
            FancyBboxPatch(
                (x - 0.72, 15.75),
                1.44,
                0.32,
                boxstyle="round,pad=0.03,rounding_size=0.06",
                facecolor=color,
                edgecolor=GRAY,
                linewidth=0.7,
            )
        )
        ax.text(x, 15.91, label, ha="center", va="center", fontsize=6.3, color=MOCHA)

    main_w = 5.7
    side_w = 2.75

    add_initial_node(ax, 5.5, 15.30)
    add_activity(ax, 5.5, 14.68, main_w, 0.44, "Khách hàng: Truy cập website", "#FFF7EF", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 14.00, main_w, 0.44, "Hệ thống Web: Hiển thị menu và Cake Builder", "#EAF4FE", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 13.32, main_w, 0.44, "Khách hàng: Xem/lọc sản phẩm hoặc tùy chỉnh bánh", "#FFF7EF", fontsize=7.0, width=36)

    add_decision(ax, 5.5, 12.55, 2.05, 0.72, "Cần AI tư vấn?")
    add_activity(ax, 8.45, 12.55, 3.05, 0.48, "Hệ thống AI: Tư vấn và gợi ý bánh", "#EDE7FF", fontsize=6.8, width=22)
    add_decision(ax, 5.5, 11.82, 0.78, 0.48, "")
    add_activity(ax, 5.5, 11.18, main_w, 0.44, "Hệ thống Web: Tính giá tự động và hiển thị tóm tắt", "#EAF4FE", fontsize=7.0, width=36)

    add_decision(ax, 5.5, 10.43, 2.25, 0.76, "Khách xác nhận đặt bánh?")
    add_activity(ax, 2.0, 10.43, side_w, 0.46, "Không tạo đơn, khách tiếp tục xem/chỉnh sửa", "#FFF7EF", fontsize=6.6, width=22)
    add_final_node(ax, 0.45, 10.43)

    add_decision(ax, 5.5, 9.68, 1.9, 0.72, "Đã đăng nhập?")
    add_activity(ax, 2.0, 9.68, side_w, 0.46, "Đăng ký / đăng nhập tài khoản", "#FFF7EF", fontsize=6.8, width=22)
    add_decision(ax, 5.5, 8.98, 0.78, 0.48, "")
    add_activity(ax, 5.5, 8.36, main_w, 0.44, "Khách hàng: Nhập thông tin nhận bánh và ngày nhận", "#FFF7EF", fontsize=7.0, width=36)

    add_decision(ax, 5.5, 7.62, 2.05, 0.72, "Thông tin hợp lệ?")
    add_activity(ax, 8.7, 7.62, 3.15, 0.46, "Hệ thống Web: Báo lỗi và yêu cầu nhập lại", "#FFEDE9", fontsize=6.6, width=24)
    add_final_node(ax, 10.55, 7.62)

    add_activity(ax, 5.5, 6.90, main_w, 0.44, "Khách hàng: Chọn thanh toán khi nhận bánh (COD)", "#FFF7EF", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 6.25, main_w, 0.44, "Hệ thống Web: Tạo đơn pending và lưu tùy chỉnh", "#EAF4FE", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 5.60, main_w, 0.44, "Hệ thống AI: Tạo phiếu đặt hàng và ghi chú cho thợ", "#EDE7FF", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 4.95, main_w, 0.44, "Admin: Xem chi tiết đơn hàng", "#FFF3D5", fontsize=7.0, width=36)

    add_decision(ax, 5.5, 4.20, 1.95, 0.72, "Đơn hợp lệ?")
    add_activity(ax, 2.0, 4.20, side_w, 0.46, "Thông báo bổ sung hoặc hủy đơn", "#FFEDE9", fontsize=6.6, width=22)
    add_final_node(ax, 0.45, 4.20)

    add_activity(ax, 5.5, 3.48, main_w, 0.44, "Admin: Xác nhận đơn pending -> confirmed", "#FFF3D5", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 2.83, main_w, 0.44, "Thợ làm bánh: Xem phiếu và cập nhật in_production", "#E7F7EF", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 2.18, main_w, 0.44, "Thợ làm bánh: Hoàn tất bánh, cập nhật ready", "#E7F7EF", fontsize=7.0, width=36)
    add_activity(ax, 5.5, 1.53, main_w, 0.44, "Khách nhận bánh, thanh toán COD; Admin cập nhật delivered", "#FFF7EF", fontsize=6.9, width=36)
    add_activity(ax, 5.5, 0.88, main_w, 0.44, "Khách hàng: Đánh giá sản phẩm", "#FFF7EF", fontsize=7.0, width=36)
    add_final_node(ax, 5.5, 0.30)

    arrow(ax, (5.5, 15.15), (5.5, 14.90))
    arrow(ax, (5.5, 14.46), (5.5, 14.22))
    arrow(ax, (5.5, 13.78), (5.5, 13.54))
    arrow(ax, (5.5, 13.10), (5.5, 12.91))
    arrow(ax, (6.53, 12.55), (6.90, 12.55), "[Có]")
    arrow(ax, (8.45, 12.31), (5.88, 11.82))
    arrow(ax, (5.5, 12.19), (5.5, 12.05), "[Không]")
    arrow(ax, (5.5, 11.58), (5.5, 11.40))
    arrow(ax, (5.5, 10.96), (5.5, 10.81))
    arrow(ax, (4.38, 10.43), (3.38, 10.43), "[Không]")
    arrow(ax, (0.64, 10.43), (0.58, 10.43))
    arrow(ax, (5.5, 10.05), (5.5, 10.04), "[Có]")
    arrow(ax, (5.5, 9.32), (5.5, 9.22), "[Rồi]")
    arrow(ax, (4.55, 9.68), (3.38, 9.68), "[Chưa]")
    arrow(ax, (2.0, 9.45), (5.05, 8.98))
    arrow(ax, (5.5, 8.74), (5.5, 8.58))
    arrow(ax, (5.5, 8.14), (5.5, 7.98))
    arrow(ax, (6.53, 7.62), (7.12, 7.62), "[Không]")
    arrow(ax, (10.28, 7.62), (10.37, 7.62))
    arrow(ax, (5.5, 7.26), (5.5, 7.12), "[Có]")
    arrow(ax, (5.5, 6.68), (5.5, 6.47))
    arrow(ax, (5.5, 6.03), (5.5, 5.82))
    arrow(ax, (5.5, 5.38), (5.5, 5.17))
    arrow(ax, (5.5, 4.73), (5.5, 4.56))
    arrow(ax, (4.53, 4.20), (3.38, 4.20), "[Không]")
    arrow(ax, (0.64, 4.20), (0.58, 4.20))
    arrow(ax, (5.5, 3.84), (5.5, 3.70), "[Có]")
    arrow(ax, (5.5, 3.26), (5.5, 3.05))
    arrow(ax, (5.5, 2.61), (5.5, 2.40))
    arrow(ax, (5.5, 1.96), (5.5, 1.75))
    arrow(ax, (5.5, 1.31), (5.5, 1.10))
    arrow(ax, (5.5, 0.66), (5.5, 0.48))

    fig.tight_layout(pad=0.35)
    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def _create_activity_diagram_v6():
    fig, ax = plt.subplots(figsize=(10.5, 16.2), dpi=220)
    ax.set_xlim(0, 10.5)
    ax.set_ylim(-0.65, 16.2)
    ax.axis("off")
    ax.set_facecolor("white")

    edge = "#111111"
    text_color = "#222222"
    line_color = "#222222"

    def uml_activity(x, y, w, h, text, fontsize=7.0, wrap_width=24):
        patch = FancyBboxPatch(
            (x - w / 2, y - h / 2),
            w,
            h,
            boxstyle="round,pad=0.05,rounding_size=0.22",
            facecolor="white",
            edgecolor=edge,
            linewidth=1.0,
            zorder=3,
        )
        ax.add_patch(patch)
        ax.text(x, y, wrapped_text(text, wrap_width), ha="center", va="center", fontsize=fontsize, color=text_color, zorder=4)
        return patch

    def uml_decision(x, y, w, h, text, fontsize=6.8):
        points = [(x, y + h / 2), (x + w / 2, y), (x, y - h / 2), (x - w / 2, y)]
        patch = Polygon(points, closed=True, facecolor="white", edgecolor=edge, linewidth=1.0, zorder=3)
        ax.add_patch(patch)
        if text:
            ax.text(x, y, wrapped_text(text, 14), ha="center", va="center", fontsize=fontsize, color=text_color, zorder=4)
        return patch

    def uml_bar(x, y, w=1.6):
        ax.add_patch(Rectangle((x - w / 2, y - 0.045), w, 0.09, facecolor="black", edgecolor="black", linewidth=0.8, zorder=5))

    def uml_start(x, y):
        ax.add_patch(Circle((x, y), 0.16, facecolor="black", edgecolor="black", linewidth=1.0, zorder=6))

    def uml_end(x, y):
        ax.add_patch(Circle((x, y), 0.20, facecolor="white", edgecolor="black", linewidth=1.0, zorder=6))
        ax.add_patch(Circle((x, y), 0.13, facecolor="black", edgecolor="black", linewidth=1.0, zorder=7))

    def uml_arrow(start, end, label=None, curve=0.0, label_shift=(0, 0.12)):
        ax.add_patch(
            FancyArrowPatch(
                start,
                end,
                arrowstyle="->",
                mutation_scale=10,
                linewidth=0.85,
                color=line_color,
                connectionstyle=f"arc3,rad={curve}",
                zorder=2,
            )
        )
        if label:
            ax.text(
                (start[0] + end[0]) / 2 + label_shift[0],
                (start[1] + end[1]) / 2 + label_shift[1],
                label,
                ha="center",
                va="center",
                fontsize=6.4,
                color=text_color,
                zorder=6,
                bbox=dict(facecolor="white", edgecolor="none", pad=0.2),
            )

    ax.text(
        5.25,
        15.78,
        "Activity Diagram - Quy trình khách hàng đặt bánh và thanh toán COD",
        ha="center",
        va="center",
        fontsize=13.5,
        color=text_color,
        fontweight="bold",
    )
    ax.text(5.25, 15.42, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=7.4, color="#555555")

    uml_start(5.25, 14.95)
    uml_activity(5.25, 14.35, 2.7, 0.48, "Khách hàng truy cập website", wrap_width=22)
    uml_activity(5.25, 13.70, 3.1, 0.50, "Hệ thống hiển thị menu và Cake Builder", wrap_width=24)
    uml_bar(5.25, 13.12, 2.5)

    uml_activity(2.8, 12.45, 2.6, 0.50, "Xem/lọc sản phẩm", wrap_width=20)
    uml_activity(7.7, 12.45, 2.95, 0.50, "Tùy chỉnh bánh theo yêu cầu", wrap_width=22)
    uml_bar(5.25, 11.76, 5.15)

    uml_decision(5.25, 11.05, 1.65, 0.68, "Cần AI tư vấn?")
    uml_activity(8.2, 11.05, 2.7, 0.50, "AI tư vấn và gợi ý bánh", wrap_width=20)
    uml_bar(5.25, 10.32, 1.55)
    uml_activity(5.25, 9.70, 3.25, 0.50, "Hệ thống tính giá tự động và hiển thị tóm tắt", wrap_width=26)

    uml_decision(5.25, 8.98, 1.85, 0.70, "Khách xác nhận đặt bánh?")
    uml_activity(2.0, 8.98, 2.6, 0.50, "Quay lại xem/lọc hoặc chỉnh sửa bánh", fontsize=6.7, wrap_width=21)

    uml_decision(5.25, 8.18, 1.55, 0.64, "Đã đăng nhập?")
    uml_activity(2.0, 8.18, 2.6, 0.50, "Đăng ký / đăng nhập tài khoản", fontsize=6.7, wrap_width=21)
    uml_bar(5.25, 7.50, 1.55)

    uml_activity(5.25, 6.88, 3.25, 0.50, "Nhập thông tin nhận bánh và ngày nhận", wrap_width=25)
    uml_decision(5.25, 6.15, 1.75, 0.66, "Thông tin hợp lệ?")
    uml_activity(8.4, 6.15, 2.7, 0.55, "Thông báo lỗi và yêu cầu nhập lại", fontsize=6.7, wrap_width=21)

    uml_activity(5.25, 5.42, 3.05, 0.50, "Chọn thanh toán khi nhận bánh (COD)", wrap_width=23)
    uml_activity(5.25, 4.78, 3.2, 0.50, "Hệ thống tạo đơn pending và lưu tùy chỉnh", wrap_width=24)
    uml_bar(5.25, 4.20, 2.5)

    uml_activity(2.8, 3.55, 2.65, 0.50, "AI tạo phiếu đặt hàng cho thợ", fontsize=6.8, wrap_width=21)
    uml_activity(7.7, 3.55, 2.65, 0.50, "Admin nhận và xem đơn mới", fontsize=6.8, wrap_width=21)
    uml_bar(5.25, 2.88, 5.15)

    uml_activity(5.25, 2.28, 2.9, 0.50, "Admin xác nhận đơn hàng", wrap_width=22)
    uml_decision(5.25, 1.58, 1.55, 0.62, "Đơn hợp lệ?")
    uml_activity(2.0, 1.58, 2.6, 0.50, "Thông báo bổ sung hoặc hủy đơn", fontsize=6.7, wrap_width=21)
    uml_activity(5.25, 0.92, 3.2, 0.48, "Thợ làm bánh cập nhật trạng thái", fontsize=6.8, wrap_width=24)
    uml_activity(5.25, 0.34, 3.2, 0.48, "Khách nhận bánh, thanh toán COD và đánh giá", fontsize=6.6, wrap_width=25)
    uml_end(5.25, -0.22)

    uml_arrow((5.25, 14.79), (5.25, 14.59))
    uml_arrow((5.25, 14.11), (5.25, 13.95))
    uml_arrow((5.25, 13.45), (5.25, 13.17))
    uml_arrow((4.55, 13.08), (2.8, 12.70))
    uml_arrow((5.95, 13.08), (7.7, 12.70))
    uml_arrow((2.8, 12.20), (4.15, 11.80))
    uml_arrow((7.7, 12.20), (6.35, 11.80))
    uml_arrow((5.25, 11.72), (5.25, 11.40))
    uml_arrow((6.08, 11.05), (6.85, 11.05), "[Có]", label_shift=(0, 0.16))
    uml_arrow((8.2, 10.80), (5.85, 10.35), curve=-0.03)
    uml_arrow((5.25, 10.71), (5.25, 10.38), "[Không]", label_shift=(0.25, 0.02))
    uml_arrow((5.25, 10.28), (5.25, 9.95))
    uml_arrow((5.25, 9.45), (5.25, 9.33))
    uml_arrow((4.33, 8.98), (3.30, 8.98), "[Không]")
    uml_arrow((2.0, 9.23), (2.72, 12.20), curve=-0.18)
    uml_arrow((5.25, 8.63), (5.25, 8.50), "[Có]")
    uml_arrow((4.48, 8.18), (3.30, 8.18), "[Chưa]")
    uml_arrow((2.0, 7.93), (4.50, 7.50), curve=0.03)
    uml_arrow((5.25, 7.86), (5.25, 7.55), "[Rồi]", label_shift=(0.22, 0.0))
    uml_arrow((5.25, 7.45), (5.25, 7.13))
    uml_arrow((5.25, 6.63), (5.25, 6.48))
    uml_arrow((6.13, 6.15), (7.05, 6.15), "[Không]")
    uml_arrow((8.4, 6.43), (6.80, 6.88), curve=0.05)
    uml_arrow((5.25, 5.82), (5.25, 5.67), "[Có]")
    uml_arrow((5.25, 5.17), (5.25, 5.03))
    uml_arrow((5.25, 4.53), (5.25, 4.25))
    uml_arrow((4.55, 4.16), (2.8, 3.80))
    uml_arrow((5.95, 4.16), (7.7, 3.80))
    uml_arrow((2.8, 3.30), (4.15, 2.92))
    uml_arrow((7.7, 3.30), (6.35, 2.92))
    uml_arrow((5.25, 2.84), (5.25, 2.53))
    uml_arrow((5.25, 2.03), (5.25, 1.89))
    uml_arrow((4.48, 1.58), (3.30, 1.58), "[Không]")
    uml_arrow((2.0, 1.33), (5.04, -0.22), curve=0.12)
    uml_arrow((5.25, 1.27), (5.25, 1.16), "[Có]")
    uml_arrow((5.25, 0.68), (5.25, 0.58))
    uml_arrow((5.25, 0.10), (5.25, -0.03))

    fig.savefig(ACTIVITY_PATH, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_activity_diagram(version: int = 6):
    """Canonical function consolidating all activity diagram styles."""
    if version == 1:
        _create_activity_diagram_v1()
    elif version == 2:
        _create_activity_diagram_v2()
    elif version == 3:
        _create_activity_diagram_v3()
    elif version == 4:
        _create_activity_diagram_v4()
    elif version == 5:
        _create_activity_diagram_v5()
    else:
        _create_activity_diagram_v6()


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 1"].font.color.rgb = RGBColor.from_string("5C3D2E")
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.color.rgb = RGBColor.from_string("5C3D2E")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 10.2, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "E8837A")
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], str(text))
    return table


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 11, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, 11)


def create_docx():
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG\nKHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(run, 12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nMÔ TẢ CHỨC NĂNG HỆ THỐNG\n")
    set_run_font(run, 16, bold=True, color="5C3D2E")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_run_font(run, 14, bold=True)

    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Tên đề tài", TOPIC),
            ("Sinh viên thực hiện", STUDENT),
            ("MSSV", MSSV),
            ("GVHD", TEACHER),
            ("Thành viên nhóm", f"{STUDENT} - thực hiện một mình"),
        ],
    )

    doc.add_heading("1. Bảng đối chiếu yêu cầu bài nộp", level=1)
    add_table(
        doc,
        ["Yêu cầu", "Phần đáp ứng trong tài liệu"],
        [
            ("Use Case không quá tổng quát", "Mục 5 và sơ đồ Mục 7 đã tách chức năng cụ thể: đăng ký, đăng nhập, đăng xuất, xem danh sách, tìm kiếm/lọc, xem chi tiết, thêm/sửa/upload/ẩn sản phẩm, lọc đơn, xác nhận đơn, cập nhật sản xuất..."),
            ("Xác định rõ Actor", "Mục 3 xác định Khách vãng lai, Khách hàng, Admin/Chủ tiệm, Thợ làm bánh và Hệ thống AI."),
            ("Activity Diagram có quy trình chính", "Mục 8 mô tả quy trình khách hàng đặt bánh và thanh toán COD."),
            ("Có điểm bắt đầu, bước xử lý, điều kiện rẽ nhánh, kết quả đầu ra, điểm kết thúc", "Activity Diagram có Bắt đầu/Kết thúc, các bước xử lý, các nhánh: cần AI tư vấn, xác nhận đặt bánh, đã đăng nhập, ngày nhận hợp lệ, admin xác nhận đơn."),
            ("File Word/PDF mô tả chức năng", "File Word này mô tả đầy đủ chức năng, actor, use case và quy trình chính."),
            ("Ảnh Use Case và ảnh quy trình chính", "Mục 7 nhúng ảnh Use Case Diagram, Mục 8 nhúng ảnh Activity Diagram."),
            ("Ghi rõ tên đề tài và thành viên", "Trang đầu và chân ảnh đều ghi đề tài, sinh viên Lý Thành Long, MSSV 2251220144."),
        ],
    )

    doc.add_heading("2. Phạm vi hệ thống và lựa chọn thanh toán", level=1)
    add_para(doc, "Hệ thống là website bán bánh kem tích hợp AI, cho phép khách xem sản phẩm, thiết kế bánh theo ý muốn, nhận tư vấn từ chatbot AI, đặt bánh, theo dõi đơn hàng và đánh giá sản phẩm. Admin quản lý sản phẩm và đơn hàng; thợ làm bánh xem phiếu sản xuất và cập nhật tiến độ.")
    add_para(doc, "Trong phạm vi bài nộp này, phương thức thanh toán được chọn là thanh toán khi nhận bánh (COD). Đây là phương án dễ triển khai, phù hợp với code hiện có và vẫn đáp ứng luồng đặt hàng - thanh toán. Các hình thức VNPay, ví điện tử, Zalo OA hoặc chuyển khoản tự động được xem là hướng phát triển mở rộng, không đưa vào chức năng bắt buộc của sơ đồ chính.", bold=True)

    doc.add_heading("3. Actor trong hệ thống", level=1)
    add_table(
        doc,
        ["Actor", "Mô tả", "Chức năng chính"],
        [
            ("Khách vãng lai", "Người chưa đăng nhập.", "Xem danh sách sản phẩm, tìm kiếm/lọc sản phẩm, xem chi tiết sản phẩm, đăng ký, đăng nhập."),
            ("Khách hàng", "Người dùng đã có tài khoản customer.", "Đăng xuất, thiết kế bánh, chat AI, xem giá tự động, đặt bánh, chọn lịch nhận, chọn COD, theo dõi đơn, đánh giá sản phẩm."),
            ("Admin / Chủ tiệm", "Người quản trị hệ thống.", "Thêm/sửa/upload ảnh/ẩn hiện sản phẩm, xem/lọc/tìm đơn, xem chi tiết đơn, xác nhận đơn, đánh dấu đã giao."),
            ("Thợ làm bánh", "Người xử lý đơn sau khi admin xác nhận.", "Xem đơn cần làm, xem phiếu làm bánh, ghi chú sản xuất, cập nhật trạng thái đang làm hoặc sẵn sàng."),
            ("Hệ thống AI", "Dịch vụ AI tích hợp vào website.", "Xử lý câu hỏi, gợi ý sản phẩm, lưu lịch sử trò chuyện, tạo tóm tắt/phiếu đặt hàng."),
        ],
    )

    doc.add_heading("4. Chức năng hệ thống theo từng nhóm", level=1)
    add_bullet(doc, "Xác thực: đăng ký tài khoản, đăng nhập, đăng xuất, kiểm tra quyền truy cập theo vai trò customer/admin/baker.")
    add_bullet(doc, "Sản phẩm phía khách: xem danh sách sản phẩm, tìm kiếm/lọc theo danh mục, xem chi tiết sản phẩm, xem giá và đánh giá.")
    add_bullet(doc, "Thiết kế bánh: chọn kích thước, vị bánh, loại kem, màu kem, topping, ghi chú đặc biệt; hệ thống cập nhật preview và tính giá tự động.")
    add_bullet(doc, "AI chatbot: người dùng đặt câu hỏi, AI xử lý dựa trên catalog, trả lời tư vấn, gợi ý sản phẩm và lưu lịch sử trò chuyện.")
    add_bullet(doc, "Đặt hàng: khách xác nhận thiết kế/sản phẩm, nhập thông tin nhận bánh, chọn ngày nhận hợp lệ, chọn COD và tạo đơn hàng trạng thái pending.")
    add_bullet(doc, "Quản trị sản phẩm: admin xem danh sách, tìm kiếm/lọc, thêm, sửa, upload ảnh, ẩn/hiện sản phẩm.")
    add_bullet(doc, "Quản trị đơn hàng: admin xem danh sách, lọc/tìm đơn, xem chi tiết, xác nhận đơn, đánh dấu đã giao.")
    add_bullet(doc, "Sản xuất: thợ làm bánh xem đơn cần làm, xem phiếu chi tiết, ghi chú sản xuất, cập nhật trạng thái confirmed → in_production → ready.")
    add_bullet(doc, "Đánh giá: khách đánh giá sản phẩm sau khi đơn hàng delivered.")

    doc.add_heading("5. Danh sách Use Case chi tiết", level=1)
    use_case_rows = [
        ("UC01", "Đăng ký tài khoản", "Khách vãng lai", "Tạo tài khoản khách hàng bằng thông tin cá nhân."),
        ("UC02", "Đăng nhập", "Khách vãng lai, Khách hàng, Admin, Thợ làm bánh", "Xác thực để sử dụng chức năng theo vai trò."),
        ("UC03", "Đăng xuất", "Khách hàng, Admin, Thợ làm bánh", "Thoát khỏi phiên đăng nhập."),
        ("UC04", "Xem danh sách sản phẩm", "Khách vãng lai, Khách hàng", "Xem menu bánh đang bán."),
        ("UC05", "Tìm kiếm / lọc sản phẩm", "Khách vãng lai, Khách hàng", "Lọc theo danh mục hoặc từ khóa."),
        ("UC06", "Xem chi tiết sản phẩm", "Khách vãng lai, Khách hàng", "Xem hình ảnh, mô tả, size, vị, giá và đánh giá."),
        ("UC07", "Thiết kế bánh Click-to-Customize", "Khách hàng", "Tùy chỉnh bánh bằng giao diện trực quan."),
        ("UC08", "Chọn kích thước, vị, kem, màu, topping", "Khách hàng", "Chọn từng thuộc tính cụ thể của bánh."),
        ("UC09", "Xem giá tự động", "Khách hàng, Hệ thống AI", "Tính và hiển thị giá theo lựa chọn."),
        ("UC10", "Chatbot AI tư vấn", "Khách hàng, Hệ thống AI", "Khách đặt câu hỏi, AI tư vấn mẫu bánh phù hợp."),
        ("UC11", "Lưu lịch sử trò chuyện", "Hệ thống AI", "Lưu lại nội dung chat để duy trì ngữ cảnh."),
        ("UC12", "Đặt bánh", "Khách hàng", "Xác nhận sản phẩm hoặc thiết kế muốn đặt."),
        ("UC13", "Chọn lịch nhận bánh", "Khách hàng", "Chọn ngày nhận hợp lệ theo quy định 24h/48h và tối đa 30 ngày."),
        ("UC14", "Chọn thanh toán khi nhận bánh (COD)", "Khách hàng", "Ghi nhận phương thức thanh toán đơn giản, dễ triển khai."),
        ("UC15", "Tạo tóm tắt / phiếu đặt hàng", "Hệ thống AI", "Tổng hợp yêu cầu và ghi chú cho admin/thợ làm bánh."),
        ("UC16", "Theo dõi trạng thái đơn", "Khách hàng", "Xem trạng thái pending, confirmed, in_production, ready, delivered."),
        ("UC17", "Đánh giá sản phẩm", "Khách hàng", "Đánh giá sau khi đơn hàng đã giao."),
        ("UC18", "Xem danh sách sản phẩm quản trị", "Admin", "Admin xem toàn bộ sản phẩm trong hệ thống."),
        ("UC19", "Thêm sản phẩm", "Admin", "Tạo sản phẩm mới với tên, mô tả, danh mục, giá, size, vị."),
        ("UC20", "Sửa sản phẩm", "Admin", "Cập nhật thông tin sản phẩm."),
        ("UC21", "Upload ảnh sản phẩm", "Admin", "Tải ảnh JPEG/PNG/WebP cho sản phẩm."),
        ("UC22", "Ẩn / hiện sản phẩm", "Admin", "Ngừng bán hoặc kích hoạt sản phẩm mà không xóa dữ liệu."),
        ("UC23", "Xem danh sách đơn hàng", "Admin", "Xem các đơn hàng trong hệ thống."),
        ("UC24", "Lọc / tìm đơn hàng", "Admin", "Lọc theo trạng thái, ngày, tên khách."),
        ("UC25", "Xem chi tiết đơn hàng", "Admin", "Xem thông tin khách, sản phẩm, tùy chỉnh, lịch sử trạng thái."),
        ("UC26", "Xác nhận đơn hàng", "Admin", "Chuyển trạng thái pending sang confirmed."),
        ("UC27", "Đánh dấu đã giao", "Admin", "Chuyển trạng thái ready sang delivered khi khách nhận bánh và thanh toán COD."),
        ("UC28", "Xem đơn cần làm", "Thợ làm bánh", "Xem đơn confirmed hoặc in_production."),
        ("UC29", "Xem phiếu làm bánh", "Thợ làm bánh", "Đọc yêu cầu tùy chỉnh và ghi chú AI."),
        ("UC30", "Ghi chú sản xuất", "Thợ làm bánh", "Thêm ghi chú nội bộ trong quá trình làm bánh."),
        ("UC31", "Cập nhật trạng thái sản xuất", "Thợ làm bánh", "Chuyển confirmed → in_production → ready."),
    ]
    add_table(doc, ["Mã", "Use Case", "Actor", "Mô tả"], use_case_rows)

    doc.add_heading("6. Quy trình hoạt động chính", level=1)
    add_para(doc, "Quy trình chính được chọn là khách hàng đặt bánh và thanh toán khi nhận bánh (COD). Quy trình này có đủ actor Khách hàng, Hệ thống AI, Admin và Thợ làm bánh.")
    add_numbered(
        doc,
        [
            "Người dùng truy cập website.",
            "Người dùng xem danh sách, tìm kiếm/lọc và xem chi tiết bánh.",
            "Khách mở Cake Builder để tùy chỉnh bánh theo kích thước, vị, kem, màu, topping.",
            "Nếu cần tư vấn, khách chat với AI; AI xử lý câu hỏi, gợi ý sản phẩm và lưu lịch sử trò chuyện.",
            "Hệ thống tính giá tự động và hiển thị tóm tắt lựa chọn.",
            "Nếu khách chưa xác nhận, khách quay lại chỉnh sửa; nếu xác nhận thì tiếp tục đặt bánh.",
            "Hệ thống kiểm tra đăng nhập; nếu chưa đăng nhập, khách đăng ký hoặc đăng nhập.",
            "Khách nhập thông tin nhận bánh và ngày giờ nhận bánh.",
            "Hệ thống kiểm tra ngày nhận: bánh thường tối thiểu 24 giờ, bánh 2 tầng tối thiểu 48 giờ, tối đa 30 ngày.",
            "Nếu ngày không hợp lệ, khách nhập lại ngày nhận.",
            "Khách chọn thanh toán khi nhận bánh (COD).",
            "Hệ thống tạo đơn hàng trạng thái pending, lưu yêu cầu tùy chỉnh và AI tạo phiếu đặt hàng.",
            "Admin kiểm tra và xác nhận đơn; nếu chưa đủ thông tin thì yêu cầu bổ sung hoặc hủy đơn.",
            "Sau khi xác nhận, đơn chuyển sang confirmed và thợ làm bánh xử lý.",
            "Thợ cập nhật trạng thái in_production và ready.",
            "Khách nhận bánh, thanh toán COD; Admin cập nhật delivered.",
            "Khách theo dõi trạng thái và đánh giá sản phẩm.",
        ],
    )

    doc.add_heading("7. Sơ đồ Use Case", level=1)
    add_para(doc, "Sơ đồ dưới đây đã tách rõ chức năng cụ thể theo từng actor, tránh các cụm quá chung như “quản lý sản phẩm” hoặc “quản lý đơn hàng”.")
    doc.add_picture(str(USECASE_PATH), width=Inches(7.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("8. Activity Diagram quy trình đặt bánh và thanh toán", level=1)
    add_para(doc, "Activity Diagram dưới đây thể hiện điểm bắt đầu, các bước xử lý, điều kiện rẽ nhánh, kết quả đầu ra là đơn hàng được xử lý đến delivered, và điểm kết thúc.")
    doc.add_picture(str(ACTIVITY_PATH), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("9. Ghi chú phạm vi so với tài liệu phân tích thiết kế", level=1)
    add_para(doc, "Tài liệu phân tích thiết kế có nhắc đến VNPay, Zalo OA, thống kê doanh thu và các hướng phát triển nâng cao. Tuy nhiên, để bài nộp Use Case/Activity Diagram bám sát chức năng hiện có và dễ bảo vệ, bản mô tả này chỉ đưa các chức năng chính đã có hoặc phù hợp với phạm vi hiện tại vào sơ đồ. Các chức năng mở rộng có thể trình bày riêng trong phần hướng phát triển nếu giảng viên hỏi thêm.")

    try:
        doc.save(DOCX_PATH)
    except PermissionError:
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        print(f"Warning: Cannot save to {DOCX_PATH} because it might be open/locked. Saving to fallback path: {fallback}")
        doc.save(fallback)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    create_usecase_diagram()
    create_activity_diagram()
    create_docx()
    print(DOCX_PATH)
    print(USECASE_PATH)
    print(ACTIVITY_PATH)


if __name__ == "__main__":
    main()
