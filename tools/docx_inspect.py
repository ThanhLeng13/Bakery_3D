from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    blocks = []
    paragraphs = []
    tables = []

    for index, block in enumerate(iter_block_items(doc), start=1):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            item = {
                "index": index,
                "type": "paragraph",
                "style": block.style.name if block.style else None,
                "text": text,
            }
            blocks.append(item)
            paragraphs.append(item)
        else:
            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            item = {
                "index": index,
                "type": "table",
                "rows": rows,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
            }
            blocks.append(item)
            tables.append(item)

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )

    inline_shapes = []
    for shape in doc.inline_shapes:
        inline_shapes.append({"width": shape.width, "height": shape.height})

    return {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "sections": sections,
        "inline_shapes": inline_shapes,
        "blocks": blocks,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("paths", nargs="+", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--write-text", action="store_true")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    for path in args.paths:
        payload = inspect_docx(path)
        output = args.output_dir / f"{path.stem}.json"
        output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        if args.write_text:
            lines = []
            for item in payload["blocks"]:
                if item["type"] == "paragraph":
                    lines.append(f"[P{item['index']}][{item['style']}] {item['text']}")
                else:
                    lines.append(f"[TABLE {item['index']}]")
                    lines.extend(" | ".join(row) for row in item["rows"])
            (args.output_dir / f"{path.stem}.txt").write_text(
                "\n".join(lines), encoding="utf-8"
            )
        print(f"{path.name}: {payload['nonempty_paragraph_count']} paragraphs, "
              f"{payload['table_count']} tables, {payload['inline_shape_count']} inline images -> {output}")


if __name__ == "__main__":
    main()
