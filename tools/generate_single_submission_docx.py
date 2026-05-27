# -*- coding: utf-8 -*-
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import generate_submission_assets_final as assets

# Configurable output directory
_out_env = os.environ.get("OUTPUT_DIR")
if not _out_env:
    for idx, arg in enumerate(sys.argv):
        if arg == "--out-dir" and idx + 1 < len(sys.argv):
            _out_env = sys.argv[idx + 1]
            break

if _out_env:
    OUT = Path(_out_env)
else:
    OUT = Path(__file__).resolve().parent.parent.parent / "tai_lieu_nop"

FINAL_DOCX = OUT / "Bai_Nop_UseCase_Activity_Web_Ban_Banh_Kem_AI_Ly_Thanh_Long.docx"
FALLBACK_DOCX = OUT / "Bai_Nop_UseCase_Activity_Web_Ban_Banh_Kem_AI_Ly_Thanh_Long_Gon_Mat.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG\nKHOA CÔNG NGHỆ THÔNG TIN")
    assets.set_run_font(run, 13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nBÀI NỘP PHÂN TÍCH CHỨC NĂNG HỆ THỐNG\n")
    assets.set_run_font(run, 16, bold=True, color="5C3D2E")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("USE CASE DIAGRAM VÀ ACTIVITY DIAGRAM")
    assets.set_run_font(run, 14, bold=True, color="5C3D2E")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\nĐề tài: {assets.TITLE}")
    assets.set_run_font(run, 14, bold=True)

    assets.add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Tên đề tài", assets.TOPIC),
            ("Sinh viên thực hiện", assets.STUDENT),
            ("MSSV", assets.MSSV),
            ("GVHD", assets.TEACHER),
            ("Thành viên nhóm", f"{assets.STUDENT} - thực hiện một mình"),
        ],
    )


def create_docx():
    OUT.mkdir(parents=True, exist_ok=True)
    assets.create_usecase_diagram()
    assets.create_activity_diagram()

    doc = Document()
    assets.style_doc(doc)
    add_cover(doc)

    add_heading(doc, "1. Mục tiêu bài nộp", level=1)
    assets.add_para(
        doc,
        "Tài liệu này mô tả chức năng hệ thống Website bán bánh kem tích hợp AI theo đúng yêu cầu: "
        "Use Case Diagram phải thể hiện rõ các chức năng cụ thể và actor sử dụng; Activity Diagram phải mô tả "
        "được quy trình chính có điểm bắt đầu, các bước xử lý, điều kiện rẽ nhánh, kết quả đầu ra và điểm kết thúc.",
    )
    assets.add_para(
        doc,
        "Phương thức thanh toán trong phạm vi bài nộp được chọn là thanh toán khi nhận bánh (COD). "
        "Đây là phương án dễ triển khai, phù hợp với hệ thống hiện tại và vẫn thể hiện đầy đủ quy trình đặt hàng - thanh toán.",
        bold=True,
    )

    add_heading(doc, "2. Đối chiếu yêu cầu", level=1)
    assets.add_table(
        doc,
        ["Yêu cầu", "Nội dung đáp ứng"],
        [
            ("Không ghi chức năng quá chung chung", "Use Case được tách thành các chức năng cụ thể như đăng ký, đăng nhập, đăng xuất, xem danh sách, tìm kiếm/lọc, xem chi tiết, thêm/sửa/upload ảnh/ẩn hiện sản phẩm, lọc đơn, xác nhận đơn, cập nhật trạng thái sản xuất."),
            ("Xác định rõ Actor", "Gồm Khách vãng lai, Khách hàng, Admin/Chủ tiệm, Thợ làm bánh và Hệ thống AI."),
            ("Có Activity Diagram cho quy trình chính", "Quy trình chính là khách hàng đặt bánh và thanh toán COD."),
            ("Activity Diagram có rẽ nhánh", "Có các điều kiện: cần AI tư vấn, khách xác nhận đặt bánh, đã đăng nhập, ngày nhận hợp lệ, admin xác nhận đơn."),
            ("Có tên đề tài và thành viên", f"Đề tài: {assets.TOPIC}; thành viên: {assets.STUDENT} ({assets.MSSV})."),
        ],
    )

    add_heading(doc, "3. Actor trong hệ thống", level=1)
    assets.add_table(
        doc,
        ["Actor", "Mô tả", "Chức năng chính"],
        [
            ("Khách vãng lai", "Người chưa đăng nhập.", "Xem danh sách sản phẩm, tìm kiếm/lọc sản phẩm, xem chi tiết sản phẩm, đăng ký, đăng nhập."),
            ("Khách hàng", "Người dùng đã đăng nhập với vai trò customer.", "Đăng xuất, thiết kế bánh, chat AI, xem giá tự động, đặt bánh, chọn lịch nhận, chọn COD, theo dõi đơn, đánh giá sản phẩm."),
            ("Admin / Chủ tiệm", "Người quản trị hệ thống.", "Thêm/sửa/upload ảnh/ẩn hiện sản phẩm, xem/lọc/tìm đơn, xem chi tiết đơn, xác nhận đơn, đánh dấu đã giao."),
            ("Thợ làm bánh", "Người xử lý đơn sau khi admin xác nhận.", "Xem đơn cần làm, xem phiếu làm bánh, ghi chú sản xuất, cập nhật trạng thái đang làm hoặc sẵn sàng."),
            ("Hệ thống AI", "Dịch vụ AI tích hợp vào website.", "Xử lý câu hỏi, gợi ý sản phẩm, lưu lịch sử trò chuyện, tạo tóm tắt/phiếu đặt hàng."),
        ],
    )

    add_heading(doc, "4. Chức năng hệ thống", level=1)
    for item in [
        "Xác thực: đăng ký tài khoản, đăng nhập, đăng xuất, kiểm tra quyền truy cập theo vai trò customer/admin/baker.",
        "Sản phẩm phía khách: xem danh sách sản phẩm, tìm kiếm/lọc theo danh mục, xem chi tiết sản phẩm, xem giá và đánh giá.",
        "Thiết kế bánh: chọn kích thước, vị bánh, loại kem, màu kem, topping, ghi chú đặc biệt; hệ thống cập nhật preview và tính giá tự động.",
        "AI chatbot: người dùng đặt câu hỏi, AI xử lý dựa trên catalog, trả lời tư vấn, gợi ý sản phẩm và lưu lịch sử trò chuyện.",
        "Đặt hàng: khách xác nhận thiết kế/sản phẩm, nhập thông tin nhận bánh, chọn ngày nhận hợp lệ, chọn COD và tạo đơn hàng trạng thái pending.",
        "Quản trị sản phẩm: admin xem danh sách, tìm kiếm/lọc, thêm, sửa, upload ảnh, ẩn/hiện sản phẩm.",
        "Quản trị đơn hàng: admin xem danh sách, lọc/tìm đơn, xem chi tiết, xác nhận đơn, đánh dấu đã giao.",
        "Sản xuất: thợ làm bánh xem đơn cần làm, xem phiếu chi tiết, ghi chú sản xuất, cập nhật trạng thái confirmed → in_production → ready.",
        "Đánh giá: khách đánh giá sản phẩm sau khi đơn hàng delivered.",
    ]:
        assets.add_bullet(doc, item)

    add_heading(doc, "5. Danh sách Use Case chi tiết", level=1)
    use_case_rows = [
        ("UC01", "Đăng ký tài khoản", "Khách vãng lai", "Tạo tài khoản khách hàng bằng thông tin cá nhân."),
        ("UC02", "Đăng nhập", "Khách vãng lai, Khách hàng, Admin, Thợ làm bánh", "Xác thực để sử dụng chức năng theo vai trò."),
        ("UC03", "Đăng xuất", "Khách hàng, Admin, Thợ làm bánh", "Thoát khỏi phiên đăng nhập."),
        ("UC04", "Xem danh sách sản phẩm", "Khách vãng lai, Khách hàng", "Xem menu bánh đang bán."),
        ("UC05", "Tìm kiếm / lọc sản phẩm", "Khách vãng lai, Khách hàng", "Lọc theo danh mục hoặc từ khóa."),
        ("UC06", "Xem chi tiết sản phẩm", "Khách vãng lai, Khách hàng", "Xem hình ảnh, mô tả, size, vị, giá và đánh giá."),
        ("UC07", "Thiết kế bánh Click-to-Customize", "Khách hàng", "Tùy chỉnh bánh bằng giao diện trực quan."),
        ("UC08", "Chọn kích thước, vị, kem, màu, topping", "Khách hàng", "Chọn từng thuộc tính cụ thể của bánh."),
        ("UC09", "Xem giá tự động", "Khách hàng, Hệ thống AI", "Tính và hiển thị giá theo lựa chọn."),
        ("UC10", "Chatbot AI tư vấn", "Khách hàng, Hệ thống AI", "Khách đặt câu hỏi, AI tư vấn mẫu bánh phù hợp."),
        ("UC11", "Lưu lịch sử trò chuyện", "Hệ thống AI", "Lưu lại nội dung chat để duy trì ngữ cảnh."),
        ("UC12", "Đặt bánh", "Khách hàng", "Xác nhận sản phẩm hoặc thiết kế muốn đặt."),
        ("UC13", "Chọn lịch nhận bánh", "Khách hàng", "Chọn ngày nhận hợp lệ theo quy định 24h/48h và tối đa 30 ngày."),
        ("UC14", "Chọn thanh toán khi nhận bánh (COD)", "Khách hàng", "Ghi nhận phương thức thanh toán."),
        ("UC15", "Tạo tóm tắt / phiếu đặt hàng", "Hệ thống AI", "Tổng hợp yêu cầu và ghi chú cho admin/thợ làm bánh."),
        ("UC16", "Theo dõi trạng thái đơn", "Khách hàng", "Xem trạng thái pending, confirmed, in_production, ready, delivered."),
        ("UC17", "Đánh giá sản phẩm", "Khách hàng", "Đánh giá sau khi đơn hàng đã giao."),
        ("UC18", "Xem danh sách sản phẩm quản trị", "Admin", "Admin xem toàn bộ sản phẩm trong hệ thống."),
        ("UC19", "Thêm sản phẩm", "Admin", "Tạo sản phẩm mới với tên, mô tả, danh mục, giá, size, vị."),
        ("UC20", "Sửa sản phẩm", "Admin", "Cập nhật thông tin sản phẩm."),
        ("UC21", "Upload ảnh sản phẩm", "Admin", "Tải ảnh JPEG/PNG/WebP cho sản phẩm."),
        ("UC22", "Ẩn / hiện sản phẩm", "Admin", "Ngừng bán hoặc kích hoạt sản phẩm mà không xóa dữ liệu."),
        ("UC23", "Xem danh sách đơn hàng", "Admin", "Xem các đơn hàng trong hệ thống."),
        ("UC24", "Lọc / tìm đơn hàng", "Admin", "Lọc theo trạng thái, ngày, tên khách."),
        ("UC25", "Xem chi tiết đơn hàng", "Admin", "Xem thông tin khách, sản phẩm, tùy chỉnh, lịch sử trạng thái."),
        ("UC26", "Xác nhận đơn hàng", "Admin", "Chuyển trạng thái pending sang confirmed."),
        ("UC27", "Đánh dấu đã giao", "Admin", "Chuyển trạng thái ready sang delivered khi khách nhận bánh và thanh toán COD."),
        ("UC28", "Xem đơn cần làm", "Thợ làm bánh", "Xem đơn confirmed hoặc in_production."),
        ("UC29", "Xem phiếu làm bánh", "Thợ làm bánh", "Đọc yêu cầu tùy chỉnh và ghi chú AI."),
        ("UC30", "Ghi chú sản xuất", "Thợ làm bánh", "Thêm ghi chú nội bộ trong quá trình làm bánh."),
        ("UC31", "Cập nhật trạng thái sản xuất", "Thợ làm bánh", "Chuyển confirmed → in_production → ready."),
    ]
    assets.add_table(doc, ["Mã", "Use Case", "Actor", "Mô tả"], use_case_rows)

    add_heading(doc, "6. Use Case Diagram", level=1)
    assets.add_para(doc, "Sơ đồ dưới đây thể hiện rõ actor và từng chức năng cụ thể trong hệ thống.")
    doc.add_picture(str(assets.USECASE_PATH), width=Inches(7.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "7. Quy trình chính: khách hàng đặt bánh và thanh toán COD", level=1)
    assets.add_numbered(
        doc,
        [
            "Người dùng truy cập website.",
            "Người dùng xem danh sách, tìm kiếm/lọc và xem chi tiết bánh.",
            "Khách mở Cake Builder để tùy chỉnh bánh theo kích thước, vị, kem, màu, topping.",
            "Nếu cần tư vấn, khách chat với AI; AI xử lý câu hỏi, gợi ý sản phẩm và lưu lịch sử trò chuyện.",
            "Hệ thống tính giá tự động và hiển thị tóm tắt lựa chọn.",
            "Nếu khách chưa xác nhận, khách quay lại chỉnh sửa; nếu xác nhận thì tiếp tục đặt bánh.",
            "Hệ thống kiểm tra đăng nhập; nếu chưa đăng nhập, khách đăng ký hoặc đăng nhập.",
            "Khách nhập thông tin nhận bánh và ngày giờ nhận bánh.",
            "Hệ thống kiểm tra ngày nhận: bánh thường tối thiểu 24 giờ, bánh 2 tầng tối thiểu 48 giờ, tối đa 30 ngày.",
            "Nếu ngày không hợp lệ, khách nhập lại ngày nhận.",
            "Khách chọn thanh toán khi nhận bánh (COD).",
            "Hệ thống tạo đơn hàng trạng thái pending, lưu yêu cầu tùy chỉnh và AI tạo phiếu đặt hàng.",
            "Admin kiểm tra và xác nhận đơn; nếu chưa đủ thông tin thì yêu cầu bổ sung hoặc hủy đơn.",
            "Sau khi xác nhận, đơn chuyển sang confirmed và thợ làm bánh xử lý.",
            "Thợ cập nhật trạng thái in_production và ready.",
            "Khách nhận bánh, thanh toán COD; Admin cập nhật delivered.",
            "Khách theo dõi trạng thái và đánh giá sản phẩm.",
        ],
    )

    add_heading(doc, "8. Activity Diagram", level=1)
    assets.add_para(doc, "Sơ đồ dưới đây có đầy đủ điểm bắt đầu, bước xử lý, điều kiện rẽ nhánh, kết quả đầu ra và điểm kết thúc.")
    doc.add_picture(str(assets.ACTIVITY_PATH), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "9. Ghi chú phạm vi", level=1)
    assets.add_para(
        doc,
        "Các chức năng như VNPay, Zalo OA, thống kê doanh thu nâng cao hoặc xuất báo cáo được xem là hướng mở rộng. "
        "Trong bản nộp này chỉ trình bày các chức năng cốt lõi và phù hợp với hệ thống hiện tại để tránh sơ đồ bị quá rộng hoặc chứa chức năng chưa triển khai.",
    )

    try:
        doc.save(FINAL_DOCX)
        return FINAL_DOCX
    except PermissionError:
        doc.save(FALLBACK_DOCX)
        return FALLBACK_DOCX


if __name__ == "__main__":
    print(create_docx())
