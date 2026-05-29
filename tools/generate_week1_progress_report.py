from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\DE_TAI")
DOCX_PATH = OUT / "Bao_Cao_Tien_Do_Thuc_Tap_Tuan_1_Ly_Thanh_Long.docx"

SCHOOL = "TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
TITLE = "BÁO CÁO TIẾN ĐỘ THỰC TẬP TUẦN 1"
TOPIC = "Xây dựng Web bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"
WEEK = "Tuần 1"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, color="000000")
    return p


def add_para(doc, text="", bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 2"].font.name = "Times New Roman"


def add_info_table(doc):
    rows = [
        ("Sinh viên thực hiện", STUDENT),
        ("Mã sinh viên", MSSV),
        ("Giáo viên hướng dẫn", TEACHER),
        ("Đề tài", TOPIC),
        ("Tuần báo cáo", WEEK),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        set_cell_shading(table.rows[i].cells[0], "EDEDED")
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_doc(doc)

    add_para(doc, SCHOOL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p = add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.runs[0], size=18, bold=True)
    doc.add_paragraph()
    add_info_table(doc)

    add_heading(doc, "1. Công việc đã thực hiện trong tuần", level=1)
    for item in [
        "Tìm hiểu yêu cầu của đề tài và xác định phạm vi hệ thống website bán bánh kem tích hợp AI.",
        "Khảo sát quy trình đặt bánh thực tế: khách xem mẫu bánh, tùy chỉnh thông tin, nhận tư vấn, đặt hàng và theo dõi trạng thái.",
        "Xác định các tác nhân chính của hệ thống gồm Khách vãng lai, Khách hàng, Admin/Chủ tiệm, Thợ làm bánh và Hệ thống AI.",
        "Liệt kê các chức năng cụ thể: đăng ký, đăng nhập, xem/tìm kiếm sản phẩm, thiết kế bánh, chatbot tư vấn, đặt bánh, quản lý đơn hàng, cập nhật trạng thái và thống kê.",
        "Phác thảo Use Case Diagram chi tiết và tách sơ đồ theo từng tác nhân để dễ theo dõi.",
        "Xây dựng Activity Diagram cho quy trình chính: khách hàng đặt bánh và thanh toán khi nhận bánh (COD).",
        "Hoàn thiện tài liệu mô tả chức năng ban đầu và chỉnh sửa theo yêu cầu nộp bài.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Kết quả đạt được", level=1)
    for item in [
        "Xác định rõ tên đề tài, mục tiêu, phạm vi và các nhóm người dùng của hệ thống.",
        "Hoàn thành danh sách chức năng ở mức cụ thể, tránh mô tả quá chung chung.",
        "Tạo được Use Case Diagram tổng quát và các sơ đồ Use Case riêng theo từng tác nhân.",
        "Tạo được Activity Diagram mô tả quy trình đặt bánh và thanh toán COD có điểm bắt đầu, bước xử lý, điều kiện rẽ nhánh và điểm kết thúc.",
        "Có bản tài liệu Word mô tả chức năng hệ thống, phục vụ việc nộp và trình bày tiến độ.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Khó khăn/vướng mắc", level=1)
    for item in [
        "Ban đầu sơ đồ Use Case còn nhiều đường nối nên nhìn rối, cần tách riêng theo từng tác nhân để dễ đọc hơn.",
        "Cần điều chỉnh Activity Diagram theo đúng ký hiệu UML như initial node, final node, decision, fork/join và luồng xử lý rõ ràng.",
        "Một số chức năng AI cần giới hạn phạm vi để phù hợp thời gian thực hiện, tránh làm đề tài quá rộng.",
        "Cần thống nhất cách trình bày tài liệu Word để đúng yêu cầu của giảng viên.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Kế hoạch thực hiện tuần tiếp theo", level=1)
    for item in [
        "Hoàn thiện tài liệu phân tích thiết kế theo cấu trúc chuẩn: đặt vấn đề, mục tiêu, phạm vi, chức năng, công nghệ và kế hoạch thực hiện.",
        "Thiết kế cơ sở dữ liệu cho các bảng chính: users, products, orders, cake_customizations, chat_history và production_notes.",
        "Thiết kế giao diện chính của hệ thống: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, Cake Builder, đặt hàng và trang quản trị.",
        "Xây dựng prototype frontend bằng Next.js và Tailwind CSS.",
        "Chuẩn bị backend FastAPI cho các chức năng tài khoản, sản phẩm và đơn hàng.",
        "Tiếp tục kiểm tra, chỉnh sửa sơ đồ Use Case và Activity Diagram nếu giảng viên góp ý thêm.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign.add_run("Sinh viên thực hiện\n\n\nLý Thành Long")
    set_run_font(run, bold=True)

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build_docx())
