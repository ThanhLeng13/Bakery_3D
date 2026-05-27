from pathlib import Path
from textwrap import wrap

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import Circle, Ellipse, FancyBboxPatch, Rectangle


OUT = Path(r"D:\DE_TAI\tai_lieu_nop")
IMAGE_DIR = OUT / "usecase_tach_tung_tac_nhan"
DOCX_PATH = OUT / "UseCase_Tach_Tung_Tac_Nhan_Web_Ban_Banh_Kem_AI.docx"
ACTIVITY_PATH = OUT / "so_do_flowchart_dat_banh_ai.png"

TITLE = "XÂY DỰNG WEB BÁN BÁNH KEM TÍCH HỢP AI"
TOPIC = "Website bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"

PINK = "#E8837A"
CREAM = "#FDF6EE"
MOCHA = "#5C3D2E"
GRAY = "#665D58"
BG = "#FFFDF9"

plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False


ACTORS = [
    {
        "name": "Khách vãng lai",
        "file": "usecase_khach_vang_lai.png",
        "color": "#FFFFFF",
        "cols": 2,
        "use_cases": [
            "Xem danh sách sản phẩm",
            "Tìm kiếm / lọc sản phẩm",
            "Xem chi tiết sản phẩm",
            "Chatbot AI tư vấn",
            "Đăng ký tài khoản",
            "Đăng nhập",
        ],
    },
    {
        "name": "Khách hàng",
        "file": "usecase_khach_hang.png",
        "color": "#FFFFFF",
        "cols": 3,
        "use_cases": [
            "Đăng xuất",
            "Xem menu, chi tiết sản phẩm",
            "Thiết kế bánh Click-to-Customize",
            "Chọn kích thước, vị, kem, màu, topping",
            "Xem giá tự động",
            "Chatbot AI tư vấn",
            "Đặt bánh",
            "Chọn lịch nhận bánh",
            "Chọn thanh toán khi nhận bánh (COD)",
            "Theo dõi trạng thái đơn",
            "Đánh giá sản phẩm",
        ],
    },
    {
        "name": "Admin / Chủ tiệm",
        "file": "usecase_admin_chu_tiem.png",
        "color": "#FFE7A8",
        "cols": 3,
        "use_cases": [
            "Xem danh sách sản phẩm quản trị",
            "Thêm sản phẩm",
            "Sửa sản phẩm",
            "Upload ảnh sản phẩm",
            "Ẩn / hiện sản phẩm",
            "Xem danh sách đơn hàng",
            "Lọc / tìm đơn hàng",
            "Xem chi tiết đơn hàng",
            "Xác nhận đơn hàng",
            "Đánh dấu đã giao",
            "Thống kê doanh thu",
            "Xuất báo cáo",
            "Quản lý tài khoản",
            "Phân quyền tài khoản",
        ],
    },
    {
        "name": "Thợ làm bánh",
        "file": "usecase_tho_lam_banh.png",
        "color": "#BEE7D2",
        "cols": 2,
        "use_cases": [
            "Xem đơn cần làm",
            "Xem phiếu làm bánh",
            "Xem lịch sản xuất",
            "Ghi chú sản xuất",
            "Cập nhật trạng thái đang làm",
            "Cập nhật trạng thái sẵn sàng",
        ],
    },
    {
        "name": "Hệ thống AI",
        "file": "usecase_he_thong_ai.png",
        "color": "#D9CCFF",
        "cols": 2,
        "use_cases": [
            "Xử lý câu hỏi tư vấn",
            "Gợi ý sản phẩm phù hợp",
            "Tư vấn phong cách bánh",
            "Kiểm tra thông tin còn thiếu",
            "Tạo tóm tắt / phiếu đặt hàng",
            "Ghi chú cho thợ làm bánh",
            "Lưu lịch sử trò chuyện",
        ],
    },
]


def wrapped_text(text: str, width: int = 18) -> str:
    return "\n".join(wrap(text, width=width, break_long_words=False))


def add_actor(ax, x, y, name):
    ax.add_patch(Circle((x, y + 0.52), 0.16, facecolor="white", edgecolor=MOCHA, linewidth=1.55, zorder=3))
    ax.plot([x, x], [y + 0.36, y - 0.24], color=MOCHA, linewidth=1.55, zorder=3)
    ax.plot([x - 0.32, x + 0.32], [y + 0.10, y + 0.10], color=MOCHA, linewidth=1.55, zorder=3)
    ax.plot([x, x - 0.28], [y - 0.24, y - 0.62], color=MOCHA, linewidth=1.55, zorder=3)
    ax.plot([x, x + 0.28], [y - 0.24, y - 0.62], color=MOCHA, linewidth=1.55, zorder=3)
    ax.text(x, y - 0.88, wrapped_text(name, 13), ha="center", va="top", fontsize=10, color=MOCHA, fontweight="bold")


def add_use_case(ax, x, y, text, color, width=18):
    ellipse = Ellipse((x, y), 2.45, 0.72, facecolor=color, edgecolor=PINK, linewidth=1.45, zorder=4)
    ax.add_patch(ellipse)
    ax.text(x, y, wrapped_text(text, width), ha="center", va="center", fontsize=7.2, color=MOCHA, zorder=5)
    return ellipse


def draw_actor_diagram(actor):
    use_cases = actor["use_cases"]
    cols = actor["cols"]
    rows = (len(use_cases) + cols - 1) // cols
    height = max(7.0, 3.3 + rows * 1.2)
    width = 12.6

    fig, ax = plt.subplots(figsize=(width, height), dpi=200)
    ax.set_xlim(0, width)
    ax.set_ylim(0, height)
    ax.axis("off")
    ax.set_facecolor(BG)

    ax.text(
        width / 2,
        height - 0.45,
        f"Use Case Diagram - {actor['name']}",
        ha="center",
        va="center",
        fontsize=15,
        color=MOCHA,
        fontweight="bold",
    )
    ax.text(width / 2, height - 0.86, f"{TITLE} | {STUDENT} ({MSSV})", ha="center", va="center", fontsize=8, color=GRAY)

    boundary_x = 2.55
    boundary_y = 0.80
    boundary_w = width - 3.05
    boundary_h = height - 2.05
    ax.add_patch(Rectangle((boundary_x, boundary_y), boundary_w, boundary_h, facecolor=CREAM, edgecolor=MOCHA, linewidth=1.55, zorder=1))
    ax.text(
        boundary_x + boundary_w / 2,
        height - 1.38,
        "Hệ thống Web bán bánh kem tích hợp AI",
        ha="center",
        va="center",
        fontsize=9,
        color=MOCHA,
        fontweight="bold",
    )

    actor_x = 1.22
    actor_y = boundary_y + boundary_h / 2
    add_actor(ax, actor_x, actor_y, actor["name"])

    left = boundary_x + 1.65
    right = boundary_x + boundary_w - 1.65
    if cols == 1:
        xs = [(left + right) / 2]
    else:
        xs = [left + i * ((right - left) / (cols - 1)) for i in range(cols)]

    top = height - 2.18
    bottom = 1.45
    if rows == 1:
        ys = [(top + bottom) / 2]
    else:
        ys = [top - i * ((top - bottom) / (rows - 1)) for i in range(rows)]

    points = []
    for index, text in enumerate(use_cases):
        row = index // cols
        col = index % cols
        x = xs[col]
        y = ys[row]
        add_use_case(ax, x, y, text, actor["color"], width=18)
        points.append((x, y))

    line_start = (actor_x + 0.36, actor_y + 0.08)
    for x, y in points:
        ax.plot([line_start[0], x - 1.22], [line_start[1], y], color=GRAY, linewidth=0.75, alpha=0.58, zorder=2)

    ax.text(
        width / 2,
        0.33,
        f"Đề tài: {TOPIC} | Thành viên: {STUDENT} ({MSSV})",
        ha="center",
        va="center",
        fontsize=8.2,
        color=MOCHA,
    )

    path = IMAGE_DIR / actor["file"]
    fig.tight_layout(pad=0.35)
    fig.savefig(path, bbox_inches="tight", facecolor=BG)
    plt.close(fig)
    return path


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(13)


def add_centered_title(doc, text, size=18):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x5C, 0x3D, 0x2E)


def build_docx(image_paths):
    doc = Document()
    style_doc(doc)

    add_centered_title(doc, "USE CASE DIAGRAM TÁCH THEO TỪNG TÁC NHÂN")
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run(f"Đề tài: {TITLE}\n").bold = True
    intro.add_run(f"Thành viên: {STUDENT} ({MSSV})")

    doc.add_paragraph(
        "Tài liệu này tách sơ đồ Use Case tổng quát thành từng sơ đồ nhỏ theo từng tác nhân. "
        "Cách trình bày này giúp nhìn rõ mỗi tác nhân sử dụng những chức năng nào và tránh nhiều đường nối chồng chéo trong một ảnh."
    )

    for index, actor in enumerate(ACTORS):
        if index:
            doc.add_page_break()
        doc.add_heading(f"{index + 1}. Tác nhân: {actor['name']}", level=1)
        doc.add_paragraph("Các chức năng chính:")
        for use_case in actor["use_cases"]:
            doc.add_paragraph(use_case, style="List Bullet")
        doc.add_paragraph("Sơ đồ Use Case riêng:")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_paths[actor["file"]]), width=Inches(6.85))

    doc.add_page_break()
    doc.add_heading("6. Activity Diagram quy trình đặt bánh và thanh toán COD", level=1)
    doc.add_paragraph(
        "Activity Diagram dưới đây mô tả quy trình chính của hệ thống: khách hàng xem/chỉnh bánh, "
        "nhận tư vấn AI nếu cần, xác nhận đơn, thanh toán COD, admin xác nhận và thợ làm bánh cập nhật trạng thái."
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ACTIVITY_PATH), width=Inches(6.85))

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        doc.save(fallback)
        return fallback


def main():
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    image_paths = {}
    for actor in ACTORS:
        image_paths[actor["file"]] = draw_actor_diagram(actor)
    output = build_docx(image_paths)
    print(output)
    for path in image_paths.values():
        print(path)


if __name__ == "__main__":
    main()
