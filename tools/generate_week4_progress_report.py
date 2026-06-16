import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_week3_progress_report import (
    add_bullet,
    add_heading,
    add_para,
    set_cell_shading,
    set_run_font,
    style_doc,
)


_env_out = os.environ.get("DE_TAI_OUTPUT_DIR")
if _env_out:
    OUT = Path(_env_out)
else:
    _tools_dir = Path(__file__).resolve().parent
    _project_root = _tools_dir.parent
    OUT = _project_root if (_project_root / "backend").exists() else Path.home()

DOCX_PATH = OUT / "Bao_Cao_Tien_Do_Thuc_Tap_Tuan_4_Ly_Thanh_Long.docx"

SCHOOL = "TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
TITLE = "BÁO CÁO TIẾN ĐỘ THỰC TẬP TUẦN 4"
TOPIC = "Xây dựng Web bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"
WEEK = "Tuần 4: 08/06/2026 - 14/06/2026"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_info_table(doc):
    rows = [
        ("Sinh viên thực hiện", STUDENT),
        ("Mã sinh viên", MSSV),
        ("Giáo viên hướng dẫn", TEACHER),
        ("Đề tài", TOPIC),
        ("Tuần báo cáo", WEEK),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        set_cell_shading(table.rows[i].cells[0], "EDEDED")
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_doc(doc)

    add_para(doc, SCHOOL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    title = add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title.runs[0], size=18, bold=True)
    doc.add_paragraph()
    add_info_table(doc)

    add_heading(doc, "1. Công việc đã thực hiện trong tuần", level=1)
    for item in [
        "Bắt đầu chuyển từ prototype giao diện sang triển khai ứng dụng full-stack bằng Next.js, Tailwind CSS và FastAPI.",
        "Xây dựng các màn hình frontend chính: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng, thanh toán, lịch sử đơn hàng, chi tiết đơn hàng, đăng nhập/đăng ký, quản lý sản phẩm, quản lý đơn hàng và khu vực thợ làm bánh.",
        "Hoàn thiện luồng Cake Builder 3D bằng Three.js/React Three Fiber, cho phép khách chọn kích thước, hương vị, loại kem, màu kem, topping, trang trí từng vùng bánh và ghi chú đặc biệt.",
        "Tích hợp Cake Builder với trang thanh toán: lưu cấu hình bánh vào localStorage, tạo đơn bánh kem theo yêu cầu và lưu customization_json để admin/thợ làm bánh xem lại chi tiết.",
        "Triển khai backend FastAPI cho các nhóm chức năng: xác thực, danh mục sản phẩm, quản lý sản phẩm admin, đơn hàng, trạng thái đơn, quản lý đơn cho admin, quản lý đơn cho thợ làm bánh, kho bánh ngọt, mua bánh ngọt, đánh giá sản phẩm và chatbot AI.",
        "Bổ sung phân quyền theo vai trò customer, admin và baker; xây dựng ProtectedRoute ở frontend và dependency kiểm tra role ở backend.",
        "Xây dựng luồng đơn hàng có state machine: pending, confirmed, in_production, ready và delivered; giới hạn quyền cập nhật trạng thái theo vai trò admin hoặc baker.",
        "Triển khai quản lý tồn kho bánh ngọt theo lô và theo chi nhánh, hỗ trợ thêm lô mới, thêm hàng loạt cho nhiều chi nhánh, ẩn/hiện lô và xem tồn kho công khai theo sản phẩm.",
        "Tích hợp chatbot AI dạng streaming SSE, có phiên chat, lịch sử chat, ngữ cảnh RAG từ danh mục sản phẩm, gợi ý sản phẩm và tóm tắt yêu cầu đặt bánh.",
        "Rà soát và sửa các lỗi tích hợp: tự tạo profile khi user Supabase/OAuth thiếu bản ghi trong bảng users, dùng service-role client cho các thao tác cần vượt RLS, xử lý bánh tự thiết kế không có product_id thật, và hiển thị đầy đủ thông tin tùy chỉnh bánh trong màn hình admin/baker.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Kết quả đạt được", level=1)
    for item in [
        "Ứng dụng đã có cấu trúc frontend/backend rõ ràng, có thể chạy theo mô hình Next.js gọi API FastAPI.",
        "Hoàn thành luồng khách hàng cơ bản: xem sản phẩm, chọn bánh ngọt vào giỏ, thiết kế bánh kem 3D, thanh toán/đặt hàng và theo dõi trạng thái đơn.",
        "Hoàn thành luồng admin: xem danh sách đơn, lọc đơn, xem chi tiết khách hàng, xác nhận đơn, cập nhật trạng thái delivered và xem đầy đủ yêu cầu tùy chỉnh bánh.",
        "Hoàn thành luồng thợ làm bánh: xem đơn cần sản xuất, xem chi tiết thiết kế bánh, ghi chú sản xuất và chuyển trạng thái confirmed sang in_production rồi ready.",
        "Hoàn thành bước đầu luồng kho bánh ngọt: quản lý lô hàng theo ngày sản xuất, hạn sử dụng, chi nhánh và số lượng còn lại; frontend có màn hình quản lý kho riêng cho baker/admin.",
        "Cake Builder đã có mô hình 3D tương tác, hỗ trợ click từng vùng bánh, đổi màu/trang trí/topping, tính giá và chuyển sang checkout.",
        "Chatbot AI đã có API phiên chat, gửi tin nhắn streaming, lưu lịch sử, lấy dữ liệu sản phẩm làm ngữ cảnh và trích xuất gợi ý/tóm tắt đơn hàng.",
        "Backend có bộ test cho auth, catalog, chat, dependencies, health và orders. Khi chạy bằng virtualenv của dự án, kết quả hiện tại là 149/150 test pass; còn 1 test chat cần đồng bộ nội dung kỳ vọng của system prompt.",
        "Các phần bảo mật và ổn định đã được cải thiện: API client tự gắn JWT, xử lý 401, auth context tự refresh token, backend có request_id logging và CORS cho lỗi 500.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Khó khăn/vướng mắc", level=1)
    for item in [
        "Khối lượng chức năng tuần 4 lớn vì phải kết nối đồng thời frontend, backend, cơ sở dữ liệu Supabase, phân quyền và luồng nghiệp vụ đặt bánh.",
        "Cake Builder có dữ liệu tùy chỉnh phức tạp theo từng vùng bánh nên cần thống nhất cách lưu JSON, cách tính giá và cách hiển thị lại cho admin/thợ làm bánh.",
        "Đơn bánh tự thiết kế không phải sản phẩm có sẵn trong catalog nên cần xử lý product_id rỗng/placeholder để tránh lỗi ràng buộc khóa ngoại.",
        "Một số thao tác backend cần service-role client để vượt RLS sau khi người dùng đã được xác thực, nếu dùng anon client sẽ dễ gặp lỗi truy cập dữ liệu.",
        "Môi trường kiểm thử cần chạy bằng virtualenv của backend; Python hệ thống còn thiếu một số dependency như python-jose và pydantic-settings.",
        "Bộ test backend còn 1 test liên quan nội dung prompt RAG chưa khớp với system prompt hiện tại, cần cập nhật lại test hoặc nội dung prompt cho thống nhất.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Kế hoạch thực hiện tuần tiếp theo", level=1)
    for item in [
        "Sửa test chat còn lỗi và xử lý cảnh báo quyền ghi .pytest_cache để bộ test backend chạy sạch.",
        "Tiếp tục kiểm thử luồng end-to-end từ khách hàng đặt bánh, admin xác nhận, thợ làm bánh sản xuất đến khi admin giao hàng.",
        "Hoàn thiện giao diện responsive cho các màn hình admin, baker, checkout, order detail và Cake Builder để sử dụng tốt trên mobile.",
        "Rà soát dữ liệu Supabase, chính sách RLS, khóa ngoại và migration để đảm bảo các luồng sản phẩm, đơn hàng, kho và đánh giá hoạt động ổn định.",
        "Tối ưu chatbot AI: chỉnh prompt, cải thiện gợi ý sản phẩm, chuẩn hóa AI_Summary và xử lý fallback khi dịch vụ AI lỗi.",
        "Chuẩn bị cấu hình triển khai: frontend trên Vercel, backend bằng Docker, kiểm tra biến môi trường và cập nhật hướng dẫn chạy/nộp báo cáo.",
        "Bổ sung ảnh chụp màn hình, mô tả API và tài liệu hướng dẫn sử dụng cho khách hàng, admin và thợ làm bánh.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign.add_run("Sinh viên thực hiện\n\n\nLý Thành Long")
    set_run_font(run, bold=True)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except (PermissionError, FileNotFoundError, OSError):
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        try:
            doc.save(fallback)
            return fallback
        except (PermissionError, FileNotFoundError, OSError):
            last_resort = Path.cwd() / DOCX_PATH.name
            doc.save(last_resort)
            return last_resort


if __name__ == "__main__":
    print(build_docx())
