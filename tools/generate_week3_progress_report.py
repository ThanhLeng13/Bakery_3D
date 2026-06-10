from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\DE_TAI")
DOCX_PATH = OUT / "Bao_Cao_Tien_Do_Thuc_Tap_Tuan_3_Ly_Thanh_Long.docx"

SCHOOL = "TRƯỜNG ĐẠI HỌC KIẾN TRÚC ĐÀ NẴNG"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
TITLE = "BÁO CÁO TIẾN ĐỘ THỰC TẬP TUẦN 3"
TOPIC = "Xây dựng Web bán bánh kem tích hợp AI"
STUDENT = "Lý Thành Long"
MSSV = "2251220144"
TEACHER = "Nguyễn Tất Phú Cường"
WEEK = "Tuần 3: 01/06/2026 - 07/06/2026"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, color="000000")
    return paragraph


def add_para(doc, text="", bold=False, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 2"].font.name = "Times New Roman"


def add_info_table(doc):
    rows = [
        ("Sinh viên thực hiện", STUDENT),
        ("Mã sinh viên", MSSV),
        ("Giáo viên hướng dẫn", TEACHER),
        ("Đề tài", TOPIC),
        ("Tuần báo cáo", WEEK),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        set_cell_shading(table.rows[i].cells[0], "EDEDED")
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style_doc(doc)

    add_para(doc, SCHOOL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, FACULTY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    title = add_para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title.runs[0], size=18, bold=True)
    doc.add_paragraph()
    add_info_table(doc)

    add_heading(doc, "1. Công việc đã thực hiện trong tuần", level=1)
    for item in [
        "Bắt đầu triển khai prototype UI cho hệ thống dựa trên các màn hình đã xác định ở tuần 2.",
        "Xây dựng bố cục giao diện cho các màn hình chính: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, thiết kế bánh tùy chỉnh và đặt bánh.",
        "Thiết kế luồng thao tác của khách hàng từ xem sản phẩm, tùy chỉnh bánh, nhận tư vấn AI, xác nhận đơn đến theo dõi trạng thái.",
        "Chuẩn bị giao diện quản trị cho Admin/Chủ tiệm gồm quản lý sản phẩm, quản lý đơn hàng và xem trạng thái xử lý đơn.",
        "Chuẩn bị giao diện cho thợ làm bánh xem phiếu đặt hàng, ghi chú sản xuất và cập nhật trạng thái làm bánh.",
        "Tạo dữ liệu mẫu cho sản phẩm bánh kem, lựa chọn kích thước, vị bánh, màu kem, topping và trạng thái đơn hàng.",
        "Rà soát lại tài liệu phân tích thiết kế để đảm bảo nội dung phù hợp với prototype UI và các chức năng đã xác định.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Kết quả đạt được", level=1)
    for item in [
        "Có danh sách màn hình chính cần triển khai cho prototype UI của hệ thống.",
        "Xác định được luồng giao diện chính dành cho khách hàng khi đặt bánh và thanh toán COD.",
        "Có định hướng bố cục giao diện cho nhóm người dùng Admin/Chủ tiệm và Thợ làm bánh.",
        "Chuẩn bị được dữ liệu mẫu phục vụ việc dựng giao diện và kiểm thử luồng đặt bánh.",
        "Tài liệu phân tích thiết kế, Use Case Diagram và Activity Diagram đã tương đối đồng bộ với phần giao diện dự kiến.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Khó khăn/vướng mắc", level=1)
    for item in [
        "Cần lựa chọn bố cục giao diện vừa đẹp, dễ nhìn, vừa phù hợp với nghiệp vụ đặt bánh kem.",
        "Prototype UI có nhiều màn hình nên cần sắp xếp luồng thao tác hợp lý để khi báo cáo không bị rối.",
        "Cần thống nhất các trạng thái đơn hàng giữa Activity Diagram, tài liệu phân tích thiết kế và giao diện.",
        "Phần chatbot AI cần trình bày ở mức vừa đủ để thể hiện chức năng tư vấn, tránh làm phạm vi tuần 3 quá rộng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Kế hoạch thực hiện tuần tiếp theo", level=1)
    for item in [
        "Hoàn thiện prototype UI và chỉnh sửa lại các màn hình theo góp ý nếu có.",
        "Bắt đầu dựng frontend bằng Next.js và Tailwind CSS cho các màn hình chính đã thiết kế.",
        "Xây dựng cấu trúc component giao diện như header, menu, card sản phẩm, form đặt bánh, trạng thái đơn hàng và dashboard quản trị.",
        "Thiết kế chi tiết API backend cho sản phẩm, đơn hàng, tài khoản và cập nhật trạng thái sản xuất.",
        "Chuẩn bị tích hợp chatbot AI ở mức cơ bản để tư vấn bánh và tạo tóm tắt đơn hàng.",
        "Tiếp tục kiểm thử luồng đặt bánh từ phía khách hàng đến Admin và Thợ làm bánh.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign.add_run("Sinh viên thực hiện\n\n\nLý Thành Long")
    set_run_font(run, bold=True)

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        fallback = DOCX_PATH.with_name(DOCX_PATH.stem + "_Moi.docx")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build_docx())
